"""
Run:  python docs/generate_user_manual.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── colour palette ─────────────────────────────────────────────────────────────
NAVY   = (0x1F, 0x38, 0x64)   # dark blue — main headings
MID    = (0x2E, 0x74, 0xB5)   # medium blue — sub-headings
STEEL  = (0x44, 0x72, 0xC4)   # lighter blue — h3
TEAL   = (0x00, 0x70, 0x5E)   # step labels / TIP text
GREEN  = (0xE2, 0xEF, 0xDA)   # TIP box fill
AMBER  = (0xFF, 0xC0, 0x00)   # NOTE accent
NOTE_F = (0xFF, 0xF2, 0xCC)   # NOTE box fill
WARN_F = (0xFF, 0xE0, 0xD0)   # WARNING fill
WARN_T = (0xC0, 0x50, 0x20)   # WARNING text
GRAY_H = (0xBF, 0xBF, 0xBF)   # table header fill
TBL_BG = (0xF2, 0xF2, 0xF2)   # table alt row fill
BODY   = (0x26, 0x26, 0x26)   # body text
MUTED  = (0x59, 0x59, 0x59)   # muted text


def set_cell_bg(cell, r, g, b):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)


def set_cell_border(cell, side="bottom", color="BFBFBF", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tbl_borders = OxmlElement("w:tcBorders")
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    sz)
    b.set(qn("w:color"), color)
    tbl_borders.append(b)
    tcPr.append(tbl_borders)


def set_table_border(table, color="BFBFBF"):
    """Add outer borders to every cell in the table."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders_el = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:color"), color)
                borders_el.append(b)
            # remove any existing tcBorders first
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcPr.append(borders_el)


# ── helpers ─────────────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading("", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*NAVY)
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def h2(doc, text):
    p = doc.add_heading("", level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*MID)
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_heading("", level=3)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*STEEL)
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*BODY)
    p.paragraph_format.space_after = Pt(4)
    return p


def step(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_after  = Pt(5)
    label = p.add_run(f"Step {number}.  ")
    label.bold = True
    label.font.color.rgb = RGBColor(*TEAL)
    label.font.size = Pt(10)
    rest = p.add_run(text)
    rest.font.size = Pt(10)
    rest.font.color.rgb = RGBColor(*BODY)
    return p


def bullet_item(doc, text, indent=0.4):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*BODY)
    p.paragraph_format.space_after = Pt(3)
    return p


def callout(doc, kind, text):
    """kind: 'TIP' | 'NOTE' | 'WARNING'"""
    cfg = {
        "TIP":     (TEAL,   GREEN,  "TIP"),
        "NOTE":    (MID,    NOTE_F, "NOTE"),
        "WARNING": (WARN_T, WARN_F, "WARNING"),
    }[kind]
    label_color, fill, label_text = cfg

    # one-cell table for the callout box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, *fill)

    # clear default paragraph and build our own
    for p_old in cell.paragraphs:
        p_old._element.getparent().remove(p_old._element)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.1)

    lbl = p.add_run(f"{label_text}:  ")
    lbl.bold = True
    lbl.font.color.rgb = RGBColor(*label_color)
    lbl.font.size = Pt(10)

    rest = p.add_run(text)
    rest.font.size = Pt(10)
    rest.font.color.rgb = RGBColor(*BODY)

    set_table_border(tbl, "D0D0D0")

    # spacer after
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl


def simple_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr_row.cells[i]
        set_cell_bg(c, *NAVY)
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        fill = TBL_BG if r_idx % 2 == 0 else (0xFF, 0xFF, 0xFF)
        for c_idx, cell_text in enumerate(row_data):
            c = row.cells[c_idx]
            set_cell_bg(c, *fill)
            p = c.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*BODY)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    set_table_border(tbl, "D0D0D0")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return tbl


def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)
    run.font.size = Pt(7)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.8)

# ── Cover Page ────────────────────────────────────────────────────────────────
cover_title = doc.add_heading("", level=0)
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("SYNTHCS")
r.font.color.rgb = RGBColor(*NAVY)
r.font.size = Pt(32)
r.font.bold = True
cover_title.paragraph_format.space_after = Pt(4)

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub1.add_run("Synthetic Dataset Generator")
rs.font.color.rgb = RGBColor(*MID)
rs.font.size = Pt(18)
rs.font.bold = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs2 = sub2.add_run("User Manual  ·  Version 1.0")
rs2.font.color.rgb = RGBColor(*MUTED)
rs2.font.size = Pt(12)

doc.add_paragraph()

inst = doc.add_paragraph()
inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
ri = inst.add_run("Gordon College — College of Computer Studies\nOlongapo City, Philippines")
ri.font.color.rgb = RGBColor(*MUTED)
ri.font.size = Pt(11)

doc.add_paragraph()

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = tag.add_run("Synthetic Data. Real Results.")
rt.italic = True
rt.font.color.rgb = RGBColor(*TEAL)
rt.font.size = Pt(11)

doc.add_page_break()

# ── Table of Contents (manual) ────────────────────────────────────────────────
h1(doc, "Table of Contents")
toc_items = [
    ("1.", "Introduction"),
    ("2.", "Getting Started"),
    ("3.", "The Dashboard"),
    ("4.", "Generating a Synthetic Dataset"),
    ("5.", "Advanced Generation Options"),
    ("6.", "Viewing and Exporting Your Dataset"),
    ("7.", "Statistical Validation"),
    ("8.", "Managing Your Datasets"),
    ("9.", "Frequently Asked Questions"),
    ("10.", "Troubleshooting"),
    ("11.", "Glossary"),
    ("12.", "Contact and Support"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    rn = p.add_run(f"{num}  ")
    rn.bold = True
    rn.font.color.rgb = RGBColor(*MID)
    rn.font.size = Pt(10)
    rt2 = p.add_run(title)
    rt2.font.size = Pt(10)
    rt2.font.color.rgb = RGBColor(*BODY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Introduction")
divider(doc)

h2(doc, "1.1  What is SynthCS?")
body(doc,
    "SynthCS is a web-based synthetic data generator designed for Computer Science students and researchers. "
    "It allows you to create realistic tabular datasets without using sensitive or hard-to-find real-world data. "
    "Whether you are building a machine learning model, conducting research, or testing an application, "
    "SynthCS generates high-quality data tailored to your needs.")
body(doc,
    "SynthCS uses CTGAN (Conditional Tabular GAN) — a deep learning model that learns statistical patterns "
    "from real data and reproduces them in new, fully synthetic rows.")

h2(doc, "1.2  What Can SynthCS Do?")
for item in [
    "Search and import reference datasets directly from Kaggle, Hugging Face, UCI, OpenML, Data.gov.ph, and PSA — without leaving the app",
    "Describe your dataset in plain English and let the AI generate the schema for you",
    "Build a custom schema manually with full control over every field",
    "Generate datasets with 1,000 to 100,000 rows",
    "Apply advanced rules: temporal patterns, relational constraints, and anomaly injection",
    "Generate multi-table relational datasets (e.g., users, orders, products)",
    "Validate the statistical quality of your synthetic data automatically",
    "Export in CSV, JSON, JSONL, SQL, or Excel format",
]:
    bullet_item(doc, item)

h2(doc, "1.3  Who Is This Manual For?")
body(doc,
    "This manual is for students and researchers using SynthCS through its web interface. "
    "No programming knowledge is required. For technical and system details, refer to the SynthCS Technical Manual.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Getting Started")
divider(doc)

h2(doc, "2.1  System Requirements")
simple_table(doc,
    ["", "Requirement"],
    [
        ["Browser",            "Google Chrome 110+, Mozilla Firefox 115+, Microsoft Edge 110+"],
        ["Internet Connection","Required"],
        ["Screen Resolution",  "1280 × 720 or higher recommended"],
        ["Account",            "Free registration required"],
    ],
    col_widths=[1.5, 4.5],
)

h2(doc, "2.2  Creating an Account")
step(doc, 1, "Open your browser and go to synthcs.site")
step(doc, 2, "Click Sign Up on the homepage.")
step(doc, 3, "Enter your email address and create a password, then click Create Account.")
step(doc, 4, "Check your email inbox for a verification link and click it to activate your account.")
callout(doc, "NOTE",
    "If you do not see the verification email, check your spam or junk folder. "
    "You can also click Resend Verification on the login page.")

h2(doc, "2.3  Logging In")
step(doc, 1, "Go to synthcs.site")
step(doc, 2, "Enter your registered email and password.")
step(doc, 3, "Click Log In.")

h2(doc, "2.4  Resetting a Forgotten Password")
step(doc, 1, "On the login page, click Forgot Password.")
step(doc, 2, "Enter your registered email address and click Send Code.")
step(doc, 3, "Check your email for a 6-digit verification code.")
step(doc, 4, "Enter the code and your new password, then click Reset Password.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. The Dashboard")
divider(doc)

body(doc,
    "After logging in, you will land on the Dashboard — your central hub in SynthCS.")

h2(doc, "3.1  What You Will See")
simple_table(doc,
    ["Section", "Description"],
    [
        ["Stats Bar",          "Shows your total generated datasets, saved schemas, and downloads"],
        ["Quick Generate",     "Load a saved schema and generate a new dataset in one click"],
        ["Recent Datasets",    "Your most recently generated datasets with quick access links"],
        ["Open Schema Builder","A shortcut button to start a new dataset"],
    ],
    col_widths=[2.0, 4.0],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — GENERATING A SYNTHETIC DATASET
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Generating a Synthetic Dataset")
divider(doc)
body(doc, "SynthCS offers four ways to generate a dataset. Choose the method that best fits your workflow.")

# ── Method 1 ──────────────────────────────────────────────────────────────────
h2(doc, "Method 1 — Smart Search  (Recommended)")
body(doc,
    "Search multiple open-source repositories directly from within SynthCS. No manual downloading required.")
step(doc, 1, "Click Schema Builder in the left sidebar.")
step(doc, 2, 'Type a keyword in the search bar (e.g., "titanic", "diabetes", "crime", "employee").')
step(doc, 3, "Click Search. SynthCS will simultaneously search Kaggle, Hugging Face, UCI, OpenML, Data.gov.ph, and PSA.")
step(doc, 4, "Browse the results. You can filter by source and sort by popularity or size.")
step(doc, 5, "Click Download on your chosen dataset. SynthCS will import it automatically and detect the schema.")
step(doc, 6, "Review the auto-detected fields and adjust any settings as needed.")
step(doc, 7, "Set the Number of Rows (1,000–100,000) and click Generate.")
callout(doc, "TIP",
    "For Philippine-focused research, filter by Data.gov.ph or PSA. "
    "For general ML datasets, try Kaggle or UCI.")

# ── Method 2 ──────────────────────────────────────────────────────────────────
h2(doc, "Method 2 — LLM Schema Generator  (AI-Assisted)")
body(doc, "Describe your dataset in plain English and the AI will build the schema automatically.")
step(doc, 1, "In Schema Builder, click LLM Schema Generator.")
step(doc, 2, "Type a description of your dataset. Be as specific as possible.")
callout(doc, "TIP",
    'Example: "A cybersecurity intrusion detection log with IP addresses, attack types, timestamps, '
    'packet counts, and severity levels for a Philippine university network."')
step(doc, 3, "Click Generate Schema. The AI will generate all fields, types, and constraints.")
step(doc, 4, "Review the generated fields. You can add, remove, or edit any field.")
step(doc, 5, "Set the Number of Rows and click Generate with CTGAN.")
callout(doc, "NOTE",
    "The LLM first generates a 200-row template, then CTGAN expands it to your target row count. "
    "This may take 2–5 minutes.")

# ── Method 3 ──────────────────────────────────────────────────────────────────
h2(doc, "Method 3 — Manual Schema Builder")
body(doc, "Build your schema from scratch with complete control over every field.")
step(doc, 1, "In Schema Builder, click + Add Field to start adding columns.")
step(doc, 2, "For each field, configure the following:")
simple_table(doc,
    ["Setting", "Description"],
    [
        ["Field Name",    "The column name in your dataset"],
        ["Data Type",     "Integer, Float, String, Boolean, Date, Email, Phone, Address, Name, UUID"],
        ["Description",   "Optional hint to help the AI generate smarter values"],
        ["Min / Max",     "Value range for numeric fields"],
        ["Allowed Values","Restrict a field to a specific list (e.g., Male, Female)"],
        ["Nullable",      "Whether the field can contain missing values"],
        ["Null Rate",     "Percentage of rows that will have a missing value"],
        ["Cardinality",   "How many distinct values to generate for categorical fields"],
    ],
    col_widths=[1.6, 4.4],
)
step(doc, 3, "Add all the fields your dataset needs.")
step(doc, 4, "Set the Number of Rows and click Generate.")

# ── Method 4 ──────────────────────────────────────────────────────────────────
h2(doc, "Method 4 — Multi-Table Generation")
body(doc, "Generate multiple related tables at once with proper foreign key relationships.")
step(doc, 1, "In Schema Builder, click Multi-Table mode.")
step(doc, 2, "Choose a preset (e.g., E-Commerce, Healthcare, Banking) or build your own table structure.")
step(doc, 3, "Configure the fields and row counts for each table.")
step(doc, 4, "Click Generate All Tables.")
body(doc,
    "SynthCS will generate all tables simultaneously and maintain referential integrity between them.",
    indent=True)

h2(doc, "4.1  Schema Presets")
body(doc,
    "SynthCS includes built-in schema presets for common use cases so you do not have to build from scratch.")
simple_table(doc,
    ["Category", "Example Presets"],
    [
        ["Mock Data",      "Student records, Employee data, Product catalog"],
        ["AI Training",    "Classification data, Regression data, Clustering data"],
        ["Cybersecurity",  "Intrusion detection logs, Network traffic, Vulnerability reports"],
        ["Stress Testing", "High-volume transaction data, Load testing datasets"],
    ],
    col_widths=[1.8, 4.2],
)
body(doc, "To use a preset: In Schema Builder, click Load Preset and select from the list.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ADVANCED GENERATION OPTIONS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Advanced Generation Options")
divider(doc)
body(doc, "Before clicking Generate, you can apply advanced rules to make your data more realistic.")

h2(doc, "5.1  Temporal Rules")
body(doc, "Add realistic time-based patterns to your date and timestamp columns.")
step(doc, 1, "In the generation settings, expand the Temporal section.")
step(doc, 2, "Toggle Enable Temporal Rules on.")
step(doc, 3, "Select the timestamp column and configure:")
for item in [
    "Date Range — the start and end dates for your data",
    "Trend — whether values increase, decrease, or stay flat over time",
    "Seasonality — repeat patterns (daily, weekly, monthly)",
]:
    bullet_item(doc, item, indent=0.6)
callout(doc, "TIP", "Example use: Simulate sales data that peaks every December.")

h2(doc, "5.2  Relational Rules")
body(doc, "Define dependencies between columns to ensure logical consistency.")
step(doc, 1, "Expand the Rules section.")
step(doc, 2, "Click + Add Rule.")
step(doc, 3, "Select the source column, the condition, and the target column.")
callout(doc, "TIP",
    'Example: "salary must be greater than 20,000 when position equals Manager"')

h2(doc, "5.3  Anomaly Injection")
body(doc,
    "Intentionally introduce outliers, noise, or missing values to simulate real-world data quality issues — "
    "useful for testing data cleaning pipelines.")
step(doc, 1, "Expand the Anomalies section.")
step(doc, 2, "Toggle Enable Anomaly Injection on.")
step(doc, 3, "Set the anomaly rate (percentage of rows affected) and anomaly type (outliers, noise, or missing values).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — VIEWING AND EXPORTING
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Viewing and Exporting Your Dataset")
divider(doc)

h2(doc, "6.1  The Data Preview Screen")
body(doc, "After generation, SynthCS automatically takes you to the Data Preview screen where you can:")
for item in [
    "Browse your generated data in a paginated table",
    "See the total number of rows and columns",
    "View validation scores at a glance",
    "Export the dataset in your preferred format",
]:
    bullet_item(doc, item)

h2(doc, "6.2  Exporting Your Dataset")
step(doc, 1, "On the Data Preview screen, click the Export dropdown.")
step(doc, 2, "Select your preferred format.")
step(doc, 3, "Click Export. Your browser will download the file automatically.")
simple_table(doc,
    ["Format", "Best For"],
    [
        ["CSV",        "Excel, Google Sheets, Python (pandas), R, SPSS"],
        ["JSON",       "Web applications and REST APIs"],
        ["JSONL",      "Machine learning pipelines (one record per line)"],
        ["SQL",        "Importing directly into a database"],
        ["Excel (.xlsx)", "Microsoft Excel with formatted sheets"],
    ],
    col_widths=[1.5, 4.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — STATISTICAL VALIDATION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Statistical Validation")
divider(doc)
body(doc,
    "SynthCS automatically evaluates the quality of your synthetic data using four statistical metrics.")

h2(doc, "7.1  Validation Metrics")
simple_table(doc,
    ["Metric", "What It Measures"],
    [
        ["Distribution Similarity",
         "How closely each column's value distribution matches the original data (Wasserstein distance)"],
        ["Correlation Preservation",
         "Whether relationships between columns are maintained in the synthetic data (Pearson correlation)"],
        ["ML Utility (TSTR)",
         "Whether a machine learning model trained on synthetic data performs well when tested on real data"],
        ["Privacy Risk",
         "Whether any synthetic row is an exact duplicate of a real row from the reference dataset"],
    ],
    col_widths=[2.2, 3.8],
)

h2(doc, "7.2  Score Interpretation")
simple_table(doc,
    ["Overall Score", "Rating", "Recommendation"],
    [
        ["80%–100%",  "Good",       "Your dataset is ready to use"],
        ["60%–79%",   "Acceptable", "Suitable for most research purposes"],
        ["Below 60%", "Low",        "Consider using a larger or cleaner reference dataset"],
    ],
    col_widths=[1.5, 1.3, 3.2],
)

h2(doc, "7.3  Viewing the Full Report")
step(doc, 1, "On the Data Preview screen, click View Full Report.")
step(doc, 2, "The full validation report shows individual scores for each metric along with visual charts.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — MANAGING YOUR DATASETS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Managing Your Datasets")
divider(doc)

h2(doc, "8.1  Downloads — Viewing Past Datasets")
body(doc, "All generated datasets are saved automatically to your account.")
step(doc, 1, "Click Downloads in the left sidebar.")
step(doc, 2, "A list of all your generated datasets appears, sorted by date.")
step(doc, 3, "Click on any dataset to open its preview and download options.")

h2(doc, "8.2  Saved Schemas")
body(doc, "Save a schema to reuse it later without rebuilding from scratch.")
for item in [
    "To save: In Schema Builder, click Save Schema, give it a name, and click Save.",
    "To load: Go to Saved Schemas in the sidebar, find your schema, and click Load.",
    "To generate from a saved schema: From the Dashboard, use the Quick Generate section to select a saved schema and generate immediately.",
]:
    bullet_item(doc, item)

h2(doc, "8.3  Deleting a Dataset")
step(doc, 1, "Go to Downloads.")
step(doc, 2, "Find the dataset you want to remove and click the trash icon.")
step(doc, 3, "Confirm the deletion.")
callout(doc, "WARNING",
    "Deleted datasets cannot be recovered. Download your file before deleting.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — FAQ
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Frequently Asked Questions")
divider(doc)

faqs = [
    ("How many rows can I generate?",
     "Between 1,000 and 100,000 rows per dataset."),
    ("How long does generation take?",
     "For the LLM and Smart Search flows using CTGAN, generation typically takes 2–5 minutes. "
     "The manual schema flow is faster, usually under 30 seconds."),
    ("Can I use SynthCS offline?",
     "No. SynthCS requires an internet connection. CTGAN training runs on our servers."),
    ("My validation score is low. What should I do?",
     "A low score can happen when the reference dataset is too small (aim for at least 500 rows) or contains many missing values. "
     "Try cleaning your reference dataset before importing, or try a different dataset."),
    ("Is my data kept private?",
     "Your data is only accessible to your account. SynthCS does not share your data with other users."),
    ("Can I use SynthCS for my thesis?",
     "Yes. SynthCS is designed for academic use. In your methodology section, describe your data as "
     "AI-generated synthetic data produced using CTGAN via SynthCS."),
    ("Can I generate Philippine-context datasets?",
     "Yes. Use the Data.gov.ph or PSA filters in Smart Search, or mention Philippine context in your LLM description. "
     "SynthCS will generate Philippine-specific names, addresses, and locations automatically."),
    ("What is the difference between the LLM method and Smart Search?",
     "The LLM method creates data from an AI-generated schema — no real reference data needed. "
     "The Smart Search method downloads real data, trains CTGAN on it, and generates synthetic data that statistically mirrors the original. "
     "Smart Search produces higher fidelity results."),
]

for question, answer in faqs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    qr = p.add_run(question)
    qr.bold = True
    qr.font.color.rgb = RGBColor(*MID)
    qr.font.size = Pt(10)
    body(doc, answer, indent=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Troubleshooting")
divider(doc)

simple_table(doc,
    ["Problem", "What to Try"],
    [
        ["Cannot log in",
         "Check your email and password. Click Forgot Password to reset via email code."],
        ["Verification email not received",
         "Check your spam folder. Click Resend Verification on the login page."],
        ["Search returns no results",
         "Try a simpler or different keyword."],
        ["Generation does not finish",
         "Wait at least 5 minutes. If it times out, try a smaller row count."],
        ["Generation fails with an error",
         "Try regenerating. If it persists, try a different reference dataset."],
        ["Dataset source gives a 404 error",
         "The specific dataset may not be available. Try a different dataset."],
        ["Download not starting",
         "Check your browser's download settings and make sure pop-ups are not blocked."],
        ["Page goes blank after generation",
         "Refresh the page and click Downloads to find your dataset."],
        ["Validation score is 0%",
         "The reference dataset may have been too small or too uniform for meaningful validation."],
    ],
    col_widths=[2.4, 3.6],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "11. Glossary")
divider(doc)

simple_table(doc,
    ["Term", "Definition"],
    [
        ["Synthetic Data",
         "Artificially generated data that statistically resembles real data but was not collected from real people or events"],
        ["Reference Dataset",
         "The real data file that SynthCS learns from to generate your synthetic dataset"],
        ["CTGAN",
         "Conditional Tabular GAN — the deep learning model that generates synthetic tabular data"],
        ["GAN",
         "Generative Adversarial Network — the type of AI architecture CTGAN is based on"],
        ["Gaussian Copula",
         "A statistical method used to preserve correlations between columns when generating data"],
        ["Schema",
         "The structure of a dataset — the list of columns, their data types, and their rules"],
        ["LLM",
         "Large Language Model — the AI used in SynthCS to generate schemas from plain English descriptions"],
        ["TSTR",
         "Train on Synthetic, Test on Real — a method for evaluating whether synthetic data is useful for ML training"],
        ["Wasserstein Distance",
         "A statistical measure of how different two value distributions are from each other"],
        ["Pearson Correlation",
         "A measure of the linear relationship between two columns"],
        ["Distribution Similarity",
         "How closely the value patterns in a synthetic column match the original"],
        ["Anomaly Injection",
         "Intentionally adding outliers, noise, or missing values to simulate real-world data quality issues"],
        ["Temporal Rules",
         "Settings that add time-based patterns (trends, seasonality) to date columns"],
        ["Relational Rules",
         "Column dependency rules that ensure logical consistency between fields"],
        ["Epoch",
         "One complete pass of the training data through the CTGAN model"],
        ["Cardinality",
         "The number of distinct values a categorical column will have"],
        ["CSV",
         "Comma-Separated Values — a standard file format for tabular data"],
        ["JSONL",
         "JSON Lines — a format where each line is a separate JSON record"],
    ],
    col_widths=[1.8, 4.2],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — CONTACT AND SUPPORT
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "12. Contact and Support")
divider(doc)
body(doc, "For technical issues, contact your system administrator or thesis adviser.")

simple_table(doc,
    ["Role", "Name"],
    [
        ["Research Adviser",  "Johniel Zar Mendoza"],
        ["Technical Adviser", "Arthur Tristan Ramos"],
        ["Thesis Adviser",    "Dr. Erlinda Cassiela Abarintos"],
        ["Institution",       "Gordon College — College of Computer Studies"],
        ["Location",          "Olongapo City, Philippines"],
        ["Website",           "synthcs.site"],
    ],
    col_widths=[2.0, 4.0],
)

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("Synthetic Data. Real Results.  ·  SynthCS — Gordon College CCS © 2026")
fr.italic = True
fr.font.color.rgb = RGBColor(*MUTED)
fr.font.size = Pt(9)

# ── Save ───────────────────────────────────────────────────────────────────────
import os
out = os.path.join(os.path.dirname(__file__), "SynthCS_User_Manual.docx")
doc.save(out)
print(f"Saved: {out}")
