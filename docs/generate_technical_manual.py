"""
Run:  python docs/generate_technical_manual.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colour palette ─────────────────────────────────────────────────────────────
NAVY   = (0x1F, 0x38, 0x64)
MID    = (0x2E, 0x74, 0xB5)
STEEL  = (0x44, 0x72, 0xC4)
TEAL   = (0x00, 0x70, 0x5E)
GREEN  = (0xE2, 0xEF, 0xDA)
NOTE_F = (0xFF, 0xF2, 0xCC)
WARN_F = (0xFF, 0xE0, 0xD0)
WARN_T = (0xC0, 0x50, 0x20)
BODY   = (0x26, 0x26, 0x26)
MUTED  = (0x59, 0x59, 0x59)
TBL_BG = (0xF2, 0xF2, 0xF2)

def set_cell_bg(cell, r, g, b):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)

def set_table_border(table, color="D0D0D0"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders_el = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), color)
                borders_el.append(b)
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcPr.append(borders_el)

def h1(doc, text):
    p = doc.add_heading("", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*NAVY)
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_heading("", level=2)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*MID)
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_heading("", level=3)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*STEEL)
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*BODY)
    p.paragraph_format.space_after = Pt(4)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    return p

def step(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(5)
    label = p.add_run(f"Step {number}.  ")
    label.bold = True
    label.font.color.rgb = RGBColor(*TEAL)
    label.font.size = Pt(10)
    rest = p.add_run(text)
    rest.font.size = Pt(10)
    rest.font.color.rgb = RGBColor(*BODY)
    return p

def bullet_item(doc, text, indent=0.4):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*BODY)
    p.paragraph_format.space_after = Pt(3)
    return p

def callout(doc, kind, text):
    cfg = {
        "TIP":     (TEAL,   GREEN,  "TIP"),
        "NOTE":    (MID,    NOTE_F, "NOTE"),
        "WARNING": (WARN_T, WARN_F, "WARNING"),
    }[kind]
    label_color, fill, label_text = cfg
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, *fill)
    for p_old in cell.paragraphs:
        p_old._element.getparent().remove(p_old._element)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.1)
    lbl = p.add_run(f"{label_text}:  ")
    lbl.bold = True
    lbl.font.color.rgb = RGBColor(*label_color)
    lbl.font.size = Pt(10)
    rest = p.add_run(text)
    rest.font.size = Pt(10)
    rest.font.color.rgb = RGBColor(*BODY)
    set_table_border(tbl, "D0D0D0")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl

def simple_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr_row.cells[i]
        set_cell_bg(c, *NAVY)
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        fill = TBL_BG if r_idx % 2 == 0 else (0xFF, 0xFF, 0xFF)
        for c_idx, cell_text in enumerate(row_data):
            c = row.cells[c_idx]
            set_cell_bg(c, *fill)
            p = c.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*BODY)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    set_table_border(tbl, "D0D0D0")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return tbl

def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)
    run.font.size = Pt(7)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.8)

# ── Cover Page ────────────────────────────────────────────────────────────────
cover = doc.add_heading("", level=0)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("SYNTHCS")
r.font.color.rgb = RGBColor(*NAVY)
r.font.size = Pt(32)
r.font.bold = True

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub1.add_run("Synthetic Dataset Generator").font.color.rgb = RGBColor(*MID)
sub1.runs[0].font.size = Pt(18)
sub1.runs[0].bold = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run("Technical Manual  ·  Version 1.0").font.color.rgb = RGBColor(*MUTED)
sub2.runs[0].font.size = Pt(12)

doc.add_paragraph()
inst = doc.add_paragraph()
inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
inst.add_run("Gordon College — College of Computer Studies\nOlongapo City, Philippines").font.color.rgb = RGBColor(*MUTED)
inst.runs[0].font.size = Pt(11)

doc.add_paragraph()
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.add_run("Synthetic Data. Real Results.").font.color.rgb = RGBColor(*TEAL)
tag.runs[0].italic = True
tag.runs[0].font.size = Pt(11)

doc.add_page_break()

# ── Table of Contents ─────────────────────────────────────────────────────────
h1(doc, "Table of Contents")
for num, title in [
    ("1.", "Introduction"),
    ("2.", "System Architecture"),
    ("3.", "Installation and Setup"),
    ("4.", "API Reference"),
    ("5.", "Database Schema"),
    ("6.", "Statistical Fidelity Validation"),
    ("7.", "CTGAN Pipeline"),
    ("8.", "LLM Integration (Claude Haiku)"),
    ("9.", "Deployment"),
    ("10.", "Appendix"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    rn = p.add_run(f"{num}  ")
    rn.bold = True
    rn.font.color.rgb = RGBColor(*MID)
    rn.font.size = Pt(10)
    p.add_run(title).font.size = Pt(10)
    p.runs[1].font.color.rgb = RGBColor(*BODY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Introduction")
divider(doc)
body(doc,
    "SynthCS is a web-based synthetic tabular dataset generator developed as a thesis project at Gordon College's "
    "College of Computer Studies. It integrates Conditional Tabular GAN (CTGAN) and Large Language Model (LLM) "
    "technologies to help Computer Science students generate realistic datasets for research and development "
    "purposes without requiring access to real-world sensitive data.")

h2(doc, "1.1  System Overview")
simple_table(doc,
    ["Field", "Value"],
    [
        ["System Name",      "SynthCS — Synthetic Dataset Generator"],
        ["Version",          "1.0"],
        ["Type",             "Web-based Application"],
        ["Node.js Backend",  "Express.js — authentication, user management, LLM proxy"],
        ["Python Backend",   "FastAPI — data processing, synthesis, dataset search"],
        ["Frontend",         "React 18 + Vite (TypeScript)"],
        ["Database",         "PostgreSQL (managed by Node.js backend)"],
        ["AI Models",        "CTGAN (ctgan package), Claude Haiku (Anthropic API)"],
        ["Frontend Host",    "Vercel"],
        ["Node.js Host",     "Railway"],
        ["Python Host",      "Hugging Face Spaces (Docker)"],
        ["Target Users",     "CS students and researchers without real-world dataset access"],
    ],
    col_widths=[1.8, 4.2],
)

h2(doc, "1.2  Purpose and Scope")
body(doc,
    "This technical manual covers the architecture, installation, configuration, API reference, and deployment "
    "procedures for SynthCS. It is intended for developers, system administrators, and technical advisers "
    "involved in deploying and maintaining the system.")

h2(doc, "1.3  Key Features")
for item in [
    "CTGAN-based synthetic tabular data generation from real reference datasets (with Gaussian Copula fallback)",
    "Smart Search — simultaneously searches Kaggle, Hugging Face, UCI, OpenML, Data.gov.ph, and PSA and downloads datasets automatically",
    "LLM-powered schema generation from plain-English descriptions (Claude Haiku)",
    "AI field suggestion and schema augmentation",
    "Manual schema builder with full field configuration",
    "Multi-table relational dataset generation with foreign key consistency",
    "Temporal rules, relational rules, and anomaly injection as post-processing options",
    "Statistical fidelity validation module (Wasserstein, Pearson, TSTR, Privacy Risk)",
    "Export in CSV, JSON, JSONL, SQL, and Excel formats",
    "PostgreSQL persistence for user accounts, saved schemas, and generated dataset metadata",
]:
    bullet_item(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. System Architecture")
divider(doc)

h2(doc, "2.1  High-Level Architecture")
body(doc,
    "SynthCS follows a three-tier architecture. The presentation layer is a React/Vite SPA deployed on Vercel. "
    "The application layer consists of two separate backends: a Node.js/Express server on Railway handling "
    "authentication, user management, and LLM calls; and a Python/FastAPI server on Hugging Face Spaces handling "
    "all data processing, dataset search, and synthetic data generation. The data layer is a PostgreSQL database "
    "on Railway, managed exclusively by the Node.js backend.")

h2(doc, "2.2  Technology Stack")
simple_table(doc,
    ["Layer", "Technology", "Purpose"],
    [
        ["Frontend",          "React 18 + Vite (TypeScript)",         "SPA user interface"],
        ["Frontend Host",     "Vercel",                               "CDN deployment and HTTPS"],
        ["Node.js Backend",   "Express.js (Node.js)",                 "Auth, user management, LLM proxy"],
        ["Node.js Host",      "Railway",                              "Node.js service deployment"],
        ["Python Backend",    "FastAPI (Python 3.11)",                "Data processing and synthesis"],
        ["Python Host",       "Hugging Face Spaces (Docker)",         "Python service deployment"],
        ["ML Engine",         "CTGAN (ctgan package)",                "Synthetic data generation"],
        ["ML Fallback",       "Gaussian Copula (numpy/scipy)",        "Fallback synthesizer when CTGAN fails"],
        ["LLM",               "Claude Haiku (Anthropic)",             "Schema generation and augmentation"],
        ["Database",          "PostgreSQL 15",                        "Persistent data storage (Railway)"],
        ["Containerization",  "Docker",                               "Python backend packaging for HF Spaces"],
        ["Validation",        "Statistical fidelity module (main.py)","Output quality assurance"],
    ],
    col_widths=[1.7, 2.3, 2.0],
)

h2(doc, "2.3  Component Overview")
body(doc, "The SynthCS system is organized into the following main components:")
for item in [
    "UI Layer — React components for schema configuration, generation controls, data preview, and result visualization",
    "Node.js API — Express server handling authentication (JWT + OAuth), email verification, saved schemas, admin dashboard, and all Anthropic Claude API calls",
    "Python API — FastAPI server exposing all data endpoints: dataset search/download, schema analysis, CTGAN generation, multi-table generation, post-processing, and validation",
    "Generation Service — CTGAN pipeline that trains on reference data and produces synthetic rows; falls back to Gaussian Copula if CTGAN fails",
    "LLM Service — Anthropic Claude Haiku integration for schema generation from prompts, AI field suggestion, schema augmentation, and search query expansion",
    "Validation Service — Statistical fidelity checks (Wasserstein distance, Pearson correlation, TSTR, Privacy Risk) via GET /api/validate/{dataset_id}",
    "Persistence Layer — PostgreSQL tables for users, datasets, saved schemas, and OAuth accounts (managed by Node.js backend)",
]:
    bullet_item(doc, item)

h2(doc, "2.4  Data Flow")
body(doc, "The typical data flow for synthetic dataset generation follows this sequence:")
for i, text in enumerate([
    "User opens Smart Search and types a keyword. The Node.js backend expands the query using Claude Haiku, then the Python backend simultaneously searches all 6 dataset sources (Kaggle, Hugging Face, UCI, OpenML, Data.gov.ph, PSA).",
    "User selects a result. The Python backend downloads the dataset automatically to temp_datasets/{uuid}/ on the server.",
    "The Python backend analyzes the downloaded CSV and returns an inferred schema (column names, types, nullable flags, sample values) to the frontend.",
    "Optionally, the Node.js backend calls Claude Haiku to augment the schema with AI-generated fields that the real dataset is missing.",
    "User reviews and adjusts the schema in the Schema Editor, sets row count, and clicks Generate.",
    "The Python backend trains CTGAN on the downloaded CSV (up to 10,000 training rows) and generates the requested number of synthetic rows. If CTGAN fails, it falls back to the Gaussian Copula synthesizer.",
    "Post-processing applies temporal rules, relational rules, and anomaly injection if configured.",
    "The fidelity validation module computes statistical similarity scores automatically.",
    "The Node.js backend saves dataset metadata to PostgreSQL. The frontend navigates to the Data Preview page where the user can inspect and download the result.",
], 1):
    step(doc, i, text)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — INSTALLATION AND SETUP
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Installation and Setup")
divider(doc)

h2(doc, "3.1  Prerequisites")
simple_table(doc,
    ["Dependency", "Version", "Notes"],
    [
        ["Python",          "3.11+",  "Required for FastAPI backend and CTGAN"],
        ["Node.js",         "18+",    "Required for Express backend and React/Vite frontend"],
        ["PostgreSQL",      "15+",    "Required for Node.js backend (user accounts, schemas, datasets)"],
        ["Anthropic API Key","—",     "Required for all LLM features (schema generation, augmentation)"],
        ["Kaggle API Key",  "—",      "Required for Kaggle dataset search and download"],
        ["Docker",          "24+",    "Required for Python backend deployment to Hugging Face Spaces"],
    ],
    col_widths=[1.8, 1.0, 3.2],
)

h2(doc, "3.2  Environment Variables")
h3(doc, "Node.js Backend (.env in backend/)")
simple_table(doc,
    ["Variable", "Description"],
    [
        ["DATABASE_URL",      "PostgreSQL connection string (e.g. postgresql://user:pass@host:5432/synthcs)"],
        ["JWT_SECRET",        "Secret key for signing JWT tokens"],
        ["ANTHROPIC_API_KEY", "Anthropic API key for Claude Haiku LLM calls"],
        ["EMAIL_USER",        "SMTP sender address for verification and reset emails"],
        ["EMAIL_PASS",        "SMTP password"],
        ["GITHUB_CLIENT_ID",  "GitHub OAuth app client ID"],
        ["GITHUB_CLIENT_SECRET", "GitHub OAuth app client secret"],
        ["GOOGLE_CLIENT_ID",  "Google OAuth client ID"],
        ["GOOGLE_CLIENT_SECRET", "Google OAuth client secret"],
        ["FRONTEND_URL",      "Frontend base URL (e.g. https://synthcs.site)"],
    ],
    col_widths=[2.2, 3.8],
)

h3(doc, "Python Backend (.env in backend/python/)")
simple_table(doc,
    ["Variable", "Description"],
    [
        ["KAGGLE_USERNAME", "Kaggle account username for dataset API access"],
        ["KAGGLE_KEY",      "Kaggle API key for dataset search and download"],
    ],
    col_widths=[1.8, 4.2],
)
callout(doc, "NOTE",
    "On Hugging Face Spaces, KAGGLE_USERNAME and KAGGLE_KEY are set as Repository Secrets "
    "under Settings → Repository secrets, not in a .env file.")

h2(doc, "3.3  Local Development Setup")
h3(doc, "Python Backend")
for i, text in enumerate([
    "Navigate to the backend/python/ directory.",
    "Create and activate a virtual environment: python -m venv venv  then  venv\\Scripts\\activate (Windows) or source venv/bin/activate (Linux/macOS).",
    "Install dependencies: pip install -r requirements.txt",
    "Create a .env file with KAGGLE_USERNAME and KAGGLE_KEY.",
    "Start the server: uvicorn main:app --reload --port 8000",
], 1):
    step(doc, i, text)

h3(doc, "Node.js Backend")
for i, text in enumerate([
    "Navigate to the backend/ directory.",
    "Install dependencies: npm install",
    "Create a .env file with the variables listed in Section 3.2.",
    "Start the server: node server.js",
], 1):
    step(doc, i, text)

h3(doc, "Frontend")
for i, text in enumerate([
    "Navigate to the project root (synthgen-clean/).",
    "Install dependencies: npm install",
    "Start the development server: npm run dev",
    "Build for production: npm run build",
], 1):
    step(doc, i, text)

callout(doc, "NOTE",
    "The frontend root is the synthgen-clean/ directory itself, not a separate frontend/ subfolder.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. API Reference")
divider(doc)

h2(doc, "4.1  Node.js Backend Endpoints")
body(doc,
    "The Node.js backend runs on port 5000 locally and is deployed on Railway. "
    "Protected endpoints require a JWT token in the Authorization: Bearer <token> header. "
    "All endpoints are prefixed with /api/ or /auth/.")

simple_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/signup",                      "Create a new user account and send verification email"],
        ["GET",  "/verify-email?token=...",       "Activate a user account after email verification"],
        ["POST", "/login",                        "Authenticate user and return JWT token"],
        ["POST", "/forgot-password",              "Send a 6-digit reset code to the user's email"],
        ["POST", "/reset-password",               "Reset password using the valid reset code"],
        ["GET",  "/auth/github",                  "Start GitHub OAuth login flow"],
        ["GET",  "/auth/google",                  "Start Google OAuth login flow"],
        ["GET",  "/api/users/:id",                "Return a user's profile data"],
        ["POST", "/api/schemas",                  "Save a schema to the database"],
        ["GET",  "/api/schemas/:userId",          "Return all saved schemas for a user"],
        ["GET",  "/api/schema/:id",               "Return a single saved schema by ID"],
        ["POST", "/api/datasets",                 "Save dataset metadata to the database"],
        ["GET",  "/api/admin/stats",              "Return system-wide statistics (admin only)"],
        ["GET",  "/api/admin/analytics",          "Return usage analytics over time (admin only)"],
        ["GET",  "/api/admin/users",              "Return paginated user list (admin only)"],
        ["POST", "/api/llm/generate-schema",      "Generate a JSON schema from a plain-English user prompt"],
        ["POST", "/api/llm/suggest-field",        "Suggest type and constraints for a single field name"],
        ["POST", "/api/llm/augment-schema",       "Add AI-generated fields missing from a real dataset's schema"],
        ["POST", "/api/llm/expand-search-query",  "Expand a search keyword into related terms for Smart Search"],
    ],
    col_widths=[0.7, 2.5, 3.0],
)

h2(doc, "4.2  Python Backend Endpoints")
body(doc,
    "The Python backend runs on port 8000 locally and is deployed on Hugging Face Spaces at port 7860. "
    "It has no authentication layer — all requests are trusted from the frontend. "
    "All endpoints are prefixed with /api/.")

simple_table(doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET",  "/",                               "Health check — returns service status"],
        ["POST", "/api/kaggle/search",              "Search Kaggle for matching datasets"],
        ["POST", "/api/kaggle/download",            "Download a Kaggle dataset and return its schema"],
        ["POST", "/api/huggingface/search",         "Search Hugging Face for matching datasets"],
        ["POST", "/api/huggingface/download",       "Download a Hugging Face dataset"],
        ["POST", "/api/uci/search",                 "Search UCI ML Repository"],
        ["POST", "/api/uci/download",               "Download a UCI dataset"],
        ["POST", "/api/openml/search",              "Search OpenML for matching datasets"],
        ["POST", "/api/openml/download",            "Download an OpenML dataset"],
        ["POST", "/api/datagov_ph/search",          "Search Data.gov.ph (Philippine government data)"],
        ["POST", "/api/datagov_ph/download",        "Download a Data.gov.ph dataset"],
        ["POST", "/api/psa/search",                 "Search Philippine Statistics Authority portal"],
        ["POST", "/api/psa/download",               "Download a PSA dataset"],
        ["POST", "/api/upload-dataset",             "Accept a user-uploaded CSV and return its schema"],
        ["POST", "/api/smart-search",               "Search all 6 sources simultaneously with expanded query terms"],
        ["POST", "/api/dataset-peek",               "Return a data preview for a specific search result"],
        ["POST", "/api/generate",                   "Train CTGAN on a downloaded dataset and generate synthetic rows"],
        ["POST", "/api/generate-from-schema",       "Generate up to 200 rows directly from a schema definition"],
        ["POST", "/api/expand-with-ctgan",          "Scale a 200-row template to a large dataset using CTGAN"],
        ["POST", "/api/generate-hybrid",            "Hybrid path: schema + real dataset combined generation"],
        ["POST", "/api/generate-multi-table",       "Generate multiple related tables with FK consistency"],
        ["GET",  "/api/preview/{dataset_id}",       "Return first N rows of a generated CSV for in-app preview"],
        ["GET",  "/api/download-template/{id}",     "Download the 200-row template in CSV, JSON, or XLSX"],
        ["GET",  "/api/download-multi/{id}",        "Download multi-table output as ZIP or XLSX"],
        ["GET",  "/api/download-entity/{id}/{tbl}", "Download a single table from multi-table output"],
        ["GET",  "/api/validate/{dataset_id}",      "Run fidelity validation and return scores"],
        ["GET",  "/api/download/{dataset_id}",      "Download final synthetic output CSV"],
    ],
    col_widths=[0.7, 2.7, 2.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Database Schema")
divider(doc)

body(doc,
    "The PostgreSQL database is managed entirely by the Node.js backend (backend/server.js). "
    "The Python backend has no database connection. Tables are created automatically on server startup "
    "using raw SQL CREATE TABLE IF NOT EXISTS statements — no migration framework is used.")

h2(doc, "5.1  users")
simple_table(doc,
    ["Column", "Type", "Description"],
    [
        ["id",                       "UUID (PK)",       "Unique user identifier, auto-generated"],
        ["full_name",                "VARCHAR(255)",    "User's full name"],
        ["first_name",               "VARCHAR(100)",    "First name (split from full_name)"],
        ["last_name",                "VARCHAR(100)",    "Last name (split from full_name)"],
        ["email",                    "VARCHAR(255)",    "User email address (unique)"],
        ["password",                 "VARCHAR(255)",    "Bcrypt-hashed password (nullable for OAuth users)"],
        ["email_verified",           "BOOLEAN",         "Whether the email address has been verified"],
        ["verification_token",       "VARCHAR(255)",    "Email verification token (cleared after use)"],
        ["verification_token_expires","TIMESTAMPTZ",    "Expiry time for the verification token"],
        ["reset_token",              "VARCHAR(255)",    "Password reset code (6-digit)"],
        ["reset_token_expires",      "TIMESTAMP",       "Expiry time for the reset token"],
        ["is_admin",                 "BOOLEAN",         "Administrator flag"],
        ["strike_count",             "INT",             "Number of content moderation strikes"],
        ["is_banned",                "BOOLEAN",         "Whether the account is permanently banned"],
        ["ban_reason",               "TEXT",            "Reason for ban (if applicable)"],
        ["created_at",               "TIMESTAMP",       "Account creation timestamp"],
    ],
    col_widths=[1.9, 1.3, 2.8],
)

h2(doc, "5.2  datasets")
simple_table(doc,
    ["Column", "Type", "Description"],
    [
        ["id",                 "UUID (PK)",      "Unique dataset identifier"],
        ["user_id",            "UUID (FK)",      "Reference to users table"],
        ["name",               "VARCHAR(255)",   "User-assigned dataset name"],
        ["kaggle_ref",         "VARCHAR(255)",   "Source dataset reference (e.g. owner/dataset-name)"],
        ["python_dataset_id",  "VARCHAR(255)",   "UUID used to locate files in temp_datasets/ on the Python server"],
        ["row_count",          "INTEGER",        "Number of rows in the generated dataset"],
        ["status",             "VARCHAR(50)",    "Generation status (ready, pending, failed)"],
        ["source",             "VARCHAR(20)",    "Origin: llm, kaggle, huggingface, uci, openml, datagov, psa, upload"],
        ["expires_at",         "TIMESTAMP",      "Expiry time (30 days after creation)"],
        ["created_at",         "TIMESTAMP",      "Generation timestamp"],
    ],
    col_widths=[1.9, 1.3, 2.8],
)

h2(doc, "5.3  schemas")
simple_table(doc,
    ["Column", "Type", "Description"],
    [
        ["id",         "UUID (PK)",     "Unique schema identifier"],
        ["user_id",    "UUID (FK)",     "Reference to users table"],
        ["name",       "VARCHAR(255)",  "User-assigned schema name"],
        ["table_name", "VARCHAR(255)",  "Optional table name for multi-table schemas"],
        ["fields",     "JSONB",         "Full schema definition as a JSON array of field objects"],
        ["created_at", "TIMESTAMP",     "Save timestamp"],
    ],
    col_widths=[1.4, 1.3, 3.3],
)

h2(doc, "5.4  oauth_accounts")
simple_table(doc,
    ["Column", "Type", "Description"],
    [
        ["id",          "UUID (PK)",     "Unique record identifier"],
        ["user_id",     "UUID (FK)",     "Reference to users table"],
        ["provider",    "VARCHAR(50)",   "OAuth provider name (github, google)"],
        ["provider_id", "VARCHAR(255)",  "User ID from the OAuth provider"],
        ["created_at",  "TIMESTAMP",     "Link creation timestamp"],
    ],
    col_widths=[1.4, 1.3, 3.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — STATISTICAL FIDELITY VALIDATION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Statistical Fidelity Validation")
divider(doc)

h2(doc, "6.1  Overview")
body(doc,
    "The fidelity validation module is implemented in backend/python/main.py (GET /api/validate/{dataset_id}). "
    "It evaluates how well the synthetic dataset preserves the statistical properties of the original reference data "
    "by computing four metrics and combining them into an overall percentage score.")

h2(doc, "6.2  Validation Metrics")
simple_table(doc,
    ["Metric", "Method", "Description"],
    [
        ["Distribution Similarity",
         "Wasserstein Distance",
         "Measures how closely each numeric column's value distribution matches the original. "
         "Score = 1 / (1 + wasserstein_distance / std). Averaged across all numeric columns."],
        ["Correlation Preservation",
         "Pearson Correlation Matrix",
         "Computes the difference between the real and synthetic correlation matrices. "
         "Score = max(0, 1 − avg_absolute_difference)."],
        ["ML Utility (TSTR)",
         "Train on Synthetic, Test on Real",
         "Trains a Random Forest classifier on synthetic data and tests it on held-out real data. "
         "Compared against a baseline trained on real data (TRTR). Score = min(1, tstr / trtr)."],
        ["Privacy Risk",
         "Exact Match Check",
         "Counts the number of synthetic rows that are exact duplicates of real rows (tuple match). "
         "Score = 1 − (matches / total_synthetic_rows)."],
    ],
    col_widths=[1.6, 1.6, 3.0],
)

h2(doc, "6.3  Overall Score Calculation")
body(doc,
    "The overall score is a weighted average of the four metrics:")
code_block(doc, "overall = (0.4 × wasserstein_score + 0.4 × correlation_score + 0.2 × utility_score) × 100")
callout(doc, "NOTE",
    "Privacy Risk is reported separately and does not factor into the weighted overall score.")

h2(doc, "6.4  Score Interpretation")
simple_table(doc,
    ["Overall Score", "Rating", "Recommendation"],
    [
        ["80% – 100%", "Good",       "Dataset is ready to use"],
        ["60% –  79%", "Acceptable", "Suitable for most research purposes"],
        ["Below 60%",  "Low",        "Consider using a larger or cleaner reference dataset (aim for 500+ rows)"],
    ],
    col_widths=[1.5, 1.3, 3.2],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — CTGAN PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. CTGAN Pipeline")
divider(doc)

h2(doc, "7.1  Overview")
body(doc,
    "The primary synthetic data generation engine is CTGAN (Conditional Tabular GAN), implemented using the "
    "standalone ctgan Python package. The pipeline is in backend/python/generator.py. "
    "When CTGAN fails (e.g. insufficient rows or memory), the system automatically falls back to "
    "a custom Gaussian Copula synthesizer in the same file.")

h2(doc, "7.2  Generation Constants")
simple_table(doc,
    ["Constant", "Value", "Description"],
    [
        ["_MAX_TRAINING_ROWS", "10,000", "Maximum rows sampled from the reference dataset for CTGAN training"],
        ["_CTGAN_EPOCHS",      "75",     "Number of CTGAN training epochs"],
        ["_CTGAN_BATCH_SIZE",  "500",    "CTGAN training batch size"],
    ],
    col_widths=[2.0, 1.0, 3.0],
)
callout(doc, "NOTE",
    "Epochs and batch size are fixed constants in generator.py, not user-configurable parameters. "
    "They were tuned to complete training within the Hugging Face Spaces request timeout.")

h2(doc, "7.3  Training Preprocessing Steps")
for item in [
    "If the reference CSV has more than 10,000 rows, a random sample of 10,000 rows is taken (random_state=42).",
    "Columns that are entirely null are dropped.",
    "Date and timestamp columns are detected and converted to ordinal integers for numeric processing.",
    "Remaining null values in object columns are filled with the string 'unknown'; numeric column nulls are filled with the column median.",
    "An 80/20 train/test split is performed. The 20% test set is saved as test_set.csv and used later by the validation module.",
    "Discrete (categorical) columns are identified as any column with dtype == object and passed to CTGAN.",
]:
    bullet_item(doc, item)

h2(doc, "7.4  Generation Flow")
body(doc, "The following steps occur when POST /api/generate is called:")
for i, text in enumerate([
    "The largest CSV file in temp_datasets/{dataset_id}/ is loaded as the reference dataset.",
    "Preprocessing is applied (see Section 7.3).",
    "CTGAN is trained: model = CTGAN(epochs=75, batch_size=500); model.fit(train_df, discrete_columns). If this raises any exception, execution falls through to the Gaussian Copula.",
    "synthetic = model.sample(row_count) produces the requested number of rows.",
    "Date columns are converted back from ordinal integers to YYYY-MM-DD strings.",
    "User-specified column changes (renames, type casts, not-null enforcement) are applied as a post-processing step.",
    "The result is saved as temp_datasets/{dataset_id}/synthetic_output.csv and the path is returned.",
], 1):
    step(doc, i, text)

h2(doc, "7.5  Gaussian Copula Fallback")
body(doc,
    "The Gaussian Copula synthesizer (_gaussian_copula_sample in generator.py) is used when CTGAN fails. "
    "It captures inter-column correlations using a Pearson correlation matrix and Cholesky decomposition, "
    "and preserves marginal distributions via empirical CDF mapping and quantile interpolation. "
    "It works with datasets as small as 200 rows and requires no GPU.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — LLM INTEGRATION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. LLM Integration (Claude Haiku)")
divider(doc)

h2(doc, "8.1  Model Configuration")
body(doc,
    "SynthCS uses Claude Haiku (Anthropic) as the LLM backbone for all AI-assisted features. "
    "All LLM calls are made from the Node.js backend (backend/server.js) using the @anthropic-ai/sdk package. "
    "The API key is stored as the ANTHROPIC_API_KEY environment variable on Railway.")
simple_table(doc,
    ["Setting", "Value"],
    [
        ["Model ID",     "claude-haiku-4-5-20251001"],
        ["Max Tokens",   "4096"],
        ["API SDK",      "@anthropic-ai/sdk (Node.js)"],
        ["Host Backend", "Node.js (Railway) — NOT the Python backend"],
    ],
    col_widths=[1.8, 4.2],
)

h2(doc, "8.2  LLM Use Cases")
body(doc, "The LLM is used in four places in the application:")

h3(doc, "Schema Generation from User Prompt  (POST /api/llm/generate-schema)")
body(doc,
    "The user types a plain-English description of their desired dataset. "
    "The Node.js backend sends the description to Claude Haiku with a structured system prompt that instructs the model "
    "to return a JSON schema with field names, data types, constraints (min/max, enum values, date ranges), "
    "and descriptions. The schema is returned to the frontend and populates the Schema Editor.")
callout(doc, "TIP",
    'Example prompt: "A cybersecurity intrusion detection log with IP addresses, attack types, timestamps, '
    'packet counts, and severity levels for a Philippine university network."')

h3(doc, "AI Field Suggestion  (POST /api/llm/suggest-field)")
body(doc,
    "When the user clicks the sparkle icon on any field in the Schema Editor, the field name is sent to Claude Haiku. "
    "The model returns a recommended data type, description, and constraints for that specific field, "
    "saving the user from manual configuration.")

h3(doc, "Schema Augmentation — Hybrid Path  (POST /api/llm/augment-schema)")
body(doc,
    "After a real dataset is downloaded and its schema is detected, Claude Haiku is asked to compare the real schema "
    "against the user's original prompt and identify missing fields. Only the missing fields are returned as "
    "AI-generated additions. These are tagged with a distinct badge in the Schema Editor so users can tell "
    "which fields came from the real data and which were AI-generated.")

h3(doc, "Search Query Expansion  (POST /api/llm/expand-search-query)")
body(doc,
    "Before searching all 6 dataset sources, the user's keyword is sent to Claude Haiku, which expands it into "
    "5–6 related terms (synonyms, domain keywords). This increases the chance of finding relevant datasets "
    "even when they are not named exactly what the user typed.")
callout(doc, "TIP",
    'Example: "student performance" → ["academic achievement", "grades", "GPA", "attendance", "test scores"]')

h2(doc, "8.3  Content Moderation")
body(doc,
    "All LLM prompts pass through a keyword-based moderation check (isInappropriatePrompt in server.js) "
    "before reaching the Anthropic API. Flagged prompts are blocked. Users receive one strike per violation; "
    "at three strikes the account is permanently banned. This check runs before any API call, so inappropriate "
    "content never reaches Claude.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Deployment")
divider(doc)

h2(doc, "9.1  Production Deployment Overview")
simple_table(doc,
    ["Service", "Platform", "URL / Notes"],
    [
        ["Frontend (React)",      "Vercel",                "synthcs.site — auto-deploys from main branch"],
        ["Node.js Backend",       "Railway",               "Express server, PostgreSQL database on Railway"],
        ["Python Backend",        "Hugging Face Spaces",   "whysoserious1012-synthcs.hf.space — Docker container"],
    ],
    col_widths=[1.8, 1.8, 2.4],
)

h2(doc, "9.2  Python Backend — Hugging Face Spaces")
body(doc,
    "The Python backend is deployed as a Docker container on Hugging Face Spaces. "
    "The repository at huggingface.co/spaces/whysoserious1012/Synthcs must contain the following files "
    "in its main branch:")
for item in [
    "Dockerfile — builds the Python 3.11 image, installs PyTorch (CPU), installs requirements.txt, exposes port 7860, runs uvicorn on 0.0.0.0:7860",
    "README.md — must include YAML frontmatter with sdk: docker so Hugging Face recognizes it as a Docker Space",
    "main.py, generator.py, and all supporting Python service files",
    "requirements.txt",
]:
    bullet_item(doc, item)
callout(doc, "NOTE",
    "Kaggle credentials (KAGGLE_USERNAME, KAGGLE_KEY) must be added as Repository Secrets in the "
    "Hugging Face Space settings, not committed to the repository.")

h2(doc, "9.3  File Storage")
body(doc,
    "The Python backend stores all temporary files on the container's local filesystem under temp_datasets/. "
    "There is no persistent external storage.")
simple_table(doc,
    ["Path", "Contents"],
    [
        ["temp_datasets/{dataset_id}/",               "Root folder for a single generation job"],
        ["temp_datasets/{dataset_id}/*.csv",           "Downloaded reference dataset CSV(s)"],
        ["temp_datasets/{dataset_id}/template.csv",   "200-row template generated from schema"],
        ["temp_datasets/{dataset_id}/test_set.csv",   "20% held-out test set for validation"],
        ["temp_datasets/{dataset_id}/synthetic_output.csv", "Final generated synthetic dataset"],
    ],
    col_widths=[2.8, 3.2],
)
callout(doc, "WARNING",
    "Files in temp_datasets/ are not persistent across Hugging Face Spaces container restarts. "
    "Users should download their generated datasets before the container is recycled.")

h2(doc, "9.4  Common Troubleshooting")
simple_table(doc,
    ["Problem", "Resolution"],
    [
        ["Anthropic API 401 error",
         "Verify ANTHROPIC_API_KEY is set correctly in the Railway environment variables."],
        ["Kaggle download 404",
         "Check that KAGGLE_USERNAME and KAGGLE_KEY are set as HF Space repository secrets. "
         "The specific dataset may also have been removed from Kaggle."],
        ["CTGAN generation times out",
         "The 5-minute frontend timeout was exceeded. Try a smaller row count or a smaller reference dataset."],
        ["Low fidelity score",
         "Ensure the reference dataset has at least 500 rows. "
         "Very uniform or very small datasets produce low Wasserstein and correlation scores."],
        ["HF Space not responding",
         "The container may be cold-starting. Wait 30–60 seconds and retry."],
        ["PostgreSQL connection refused",
         "Check DATABASE_URL in the Railway Node.js environment variables."],
        ["Frontend 404 on page refresh",
         "Vercel is configured to serve index.html for all routes via vercel.json rewrites."],
    ],
    col_widths=[2.2, 3.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — APPENDIX
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Appendix")
divider(doc)

h2(doc, "10.1  Glossary")
simple_table(doc,
    ["Term", "Definition"],
    [
        ["CTGAN",
         "Conditional Tabular GAN — a GAN-based model specialized for generating synthetic tabular data. "
         "Implemented via the ctgan Python package."],
        ["Gaussian Copula",
         "A statistical synthesizer that captures inter-column correlations via a Pearson correlation matrix "
         "and Cholesky decomposition. Used as the fallback when CTGAN fails."],
        ["LLM",
         "Large Language Model — Claude Haiku (Anthropic) is used for schema generation, field suggestion, "
         "schema augmentation, and query expansion."],
        ["TSTR",
         "Train on Synthetic, Test on Real — evaluates whether a model trained on synthetic data performs "
         "well on real held-out data."],
        ["Wasserstein Distance",
         "A statistical measure of how different two probability distributions are from each other. "
         "Used for column-level distribution similarity scoring."],
        ["Pearson Correlation",
         "Measures the linear relationship between two columns. Used to assess correlation preservation "
         "between real and synthetic datasets."],
        ["Fidelity Score",
         "A composite metric (weighted average of Wasserstein and correlation scores plus TSTR) "
         "measuring how statistically similar synthetic data is to the reference dataset."],
        ["Smart Search",
         "A built-in feature that simultaneously searches Kaggle, Hugging Face, UCI, OpenML, "
         "Data.gov.ph, and PSA using expanded query terms."],
        ["Schema Augmentation",
         "The process of adding AI-generated fields to a real dataset's schema when the real data is "
         "missing columns requested in the user's prompt."],
        ["temp_datasets/",
         "The temporary file storage directory on the Python backend server where downloaded, "
         "template, and generated CSV files are stored per job."],
        ["FastAPI",
         "A modern Python web framework for building REST APIs. Used for the Python backend."],
        ["HF Spaces",
         "Hugging Face Spaces — the cloud hosting platform for the Python backend Docker container."],
    ],
    col_widths=[1.8, 4.2],
)

h2(doc, "10.2  References")
for ref in [
    "Xu, L. et al. (2019). Modeling Tabular Data using Conditional GAN. NeurIPS 2019.",
    "Anthropic. (2024). Claude API Documentation. https://docs.anthropic.com",
    "CTGAN Python Package. https://github.com/sdv-dev/CTGAN",
    "FastAPI Documentation. https://fastapi.tiangolo.com",
    "React Documentation. https://react.dev",
    "Hugging Face Spaces Documentation. https://huggingface.co/docs/hub/spaces",
]:
    bullet_item(doc, ref)

h2(doc, "10.3  Thesis Advisers")
simple_table(doc,
    ["Role", "Name"],
    [
        ["Research Adviser",  "Johniel Zar Mendoza"],
        ["Technical Adviser", "Arthur Tristan Ramos"],
        ["Thesis Adviser",    "Dr. Erlinda Cassiela Abarintos"],
        ["Institution",       "Gordon College — College of Computer Studies"],
        ["Location",          "Olongapo City, Philippines"],
        ["Website",           "synthcs.site"],
    ],
    col_widths=[2.0, 4.0],
)

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("Synthetic Data. Real Results.  ·  SynthCS — Gordon College CCS © 2026")
fr.italic = True
fr.font.color.rgb = RGBColor(*MUTED)
fr.font.size = Pt(9)

# ── Save ───────────────────────────────────────────────────────────────────────
import os
out = os.path.join(os.path.dirname(__file__), "SynthCS_Technical_Manual.docx")
doc.save(out)
print(f"Saved: {out}")
