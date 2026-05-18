"""
Run this script to generate the backend reviewer docx:
  python docs/generate_backend_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold_start=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    return p

def add_tag(doc, label, color_rgb):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"  {label}  ")
    run.bold = True
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

def add_divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
    p.runs[0].font.size = Pt(8)

def section(doc, eng, tag):
    """English + Tagalog side-by-side as two paragraphs."""
    p = doc.add_paragraph()
    r1 = p.add_run("🇺🇸 EN: ")
    r1.bold = True; r1.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p.add_run(eng)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0)
    r2 = p2.add_run("🇵🇭 TL: ")
    r2.bold = True; r2.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p2.add_run(tag)
    doc.add_paragraph()

def bullet(doc, eng, tag, indent=True):
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run("EN: ")
    r.bold = True; r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p.add_run(eng)
    p2 = doc.add_paragraph(style="List Bullet")
    if indent:
        p2.paragraph_format.left_indent = Inches(0.6)
    r2 = p2.add_run("TL: ")
    r2.bold = True; r2.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p2.add_run(tag)

# ── Build document ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section_obj in doc.sections:
    section_obj.top_margin    = Cm(2.0)
    section_obj.bottom_margin = Cm(2.0)
    section_obj.left_margin   = Cm(2.5)
    section_obj.right_margin  = Cm(2.5)

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("SynthCS — Backend Technical Reviewer", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)

sub = doc.add_paragraph("Comprehensive guide to the backend codebase, API endpoints, algorithms, LLM integration, and Gaussian Copula synthesizer.")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

sub2 = doc.add_paragraph("Kasama ang Tagalog na paliwanag para sa mas madaling pag-unawa.")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.runs[0].font.italic = True
sub2.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. System Architecture Overview", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

section(doc,
    "SynthCS uses a THREE-LAYER architecture: (1) React frontend deployed on Vercel, "
    "(2) Node.js backend (Express) deployed on Railway, and (3) Python backend (FastAPI) also deployed on Railway. "
    "Each layer has a specific responsibility.",
    "Ang SynthCS ay gumagamit ng TATLONG LAYER na arkitektura: (1) React frontend na naka-deploy sa Vercel, "
    "(2) Node.js backend (Express) na naka-deploy sa Railway, at (3) Python backend (FastAPI) na naka-deploy din sa Railway. "
    "Bawat layer ay may sariling tungkulin."
)

add_heading(doc, "Layer Breakdown", 2)
bullet(doc,
    "FRONTEND (React + Vercel): The user interface. Runs in the browser. "
    "Sends requests to both the Node.js and Python backends.",
    "FRONTEND (React + Vercel): Ang user interface. Tumatakbo sa browser. "
    "Nagpapadala ng mga kahilingan (requests) sa Node.js at Python backends."
)
bullet(doc,
    "NODE.JS BACKEND (server.js on Railway): Handles user accounts, authentication, "
    "email verification, saved schemas, admin dashboard, and ALL LLM (AI) calls to Anthropic's Claude.",
    "NODE.JS BACKEND (server.js sa Railway): Nangangasiwa ng user accounts, authentication, "
    "email verification, saved schemas, admin dashboard, at LAHAT ng LLM (AI) calls sa Anthropic Claude."
)
bullet(doc,
    "PYTHON BACKEND (FastAPI on Railway): Handles ALL data processing — searching datasets, "
    "downloading real datasets, generating synthetic data using the Gaussian Copula algorithm, "
    "schema-based generation, and multi-table generation.",
    "PYTHON BACKEND (FastAPI sa Railway): Nangangasiwa ng LAHAT ng data processing — "
    "paghahanap ng datasets, pag-download ng real datasets, pag-generate ng synthetic data gamit ang "
    "Gaussian Copula algorithm, schema-based generation, at multi-table generation."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Key Backend Files and Their Purpose", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

files = [
    ("backend/server.js", "Main Node.js server. Handles auth, user management, LLM calls, and admin.",
     "Pangunahing Node.js server. Nangangasiwa ng auth, user management, LLM calls, at admin."),
    ("backend/python/main.py", "Main Python FastAPI service. Defines all Python API endpoints and routes requests.",
     "Pangunahing Python FastAPI service. Nagtatakda ng lahat ng Python API endpoints."),
    ("backend/python/generator.py", "The Gaussian Copula synthesizer algorithm. Core synthetic data generation engine.",
     "Ang Gaussian Copula synthesizer algorithm. Core engine ng synthetic data generation."),
    ("backend/python/smart_gen_data.py", "Schema-based data generator. Produces realistic values column-by-column from a schema definition.",
     "Schema-based data generator. Gumagawa ng realistic na values per column mula sa schema definition."),
    ("backend/python/relational_gen.py", "Handles multi-table generation with FK consistency, entity grouping, and HR payroll logic.",
     "Nangangasiwa ng multi-table generation na may FK consistency, entity grouping, at HR payroll logic."),
    ("backend/python/analyzer.py", "Analyzes a CSV file and infers its schema (column names, types, nullable flags, sample values).",
     "Nag-a-analyze ng CSV file at nag-iinfer ng schema nito (column names, types, nullable flags, sample values)."),
    ("backend/python/kaggle_service.py", "Searches and downloads datasets from Kaggle using the Kaggle API.",
     "Naghahanap at nagda-download ng datasets mula sa Kaggle gamit ang Kaggle API."),
    ("backend/python/huggingface_service.py", "Searches and downloads datasets from HuggingFace.",
     "Naghahanap at nagda-download ng datasets mula sa HuggingFace."),
    ("backend/python/uci_service.py", "Searches and downloads datasets from the UCI ML Repository.",
     "Naghahanap at nagda-download ng datasets mula sa UCI ML Repository."),
    ("backend/python/openml_service.py", "Searches and downloads datasets from OpenML.",
     "Naghahanap at nagda-download ng datasets mula sa OpenML."),
    ("backend/python/datagov_ph_service.py", "Searches and downloads datasets from Data.gov.ph.",
     "Naghahanap at nagda-download ng datasets mula sa Data.gov.ph."),
    ("backend/python/psa_service.py", "Searches and downloads datasets from the Philippine Statistics Authority (PSA).",
     "Naghahanap at nagda-download ng datasets mula sa Philippine Statistics Authority (PSA)."),
    ("backend/python/temporal_engine.py", "Post-processing: applies temporal rules (date ranges, business hours, ordered timestamps).",
     "Post-processing: nag-a-apply ng temporal rules (date ranges, business hours, ordered timestamps)."),
    ("backend/python/relationship_engine.py", "Post-processing: enforces IF-THEN rules between columns.",
     "Post-processing: nagpapatupad ng IF-THEN rules sa pagitan ng mga columns."),
    ("backend/python/anomaly_injector.py", "Post-processing: injects anomalous records (nulls, outliers, duplicates) at a set rate.",
     "Post-processing: nag-iinject ng anomalous records (nulls, outliers, duplicates) sa itinakdang rate."),
]

for fname, eng_desc, tag_desc in files:
    p = doc.add_paragraph()
    r = p.add_run(fname)
    r.bold = True; r.font.name = "Courier New"; r.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    p.add_run(f"\n  EN: {eng_desc}")
    p.add_run(f"\n  TL: {tag_desc}")
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — API ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Complete API Endpoint Reference", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

add_heading(doc, "3.1 Node.js Endpoints (backend/server.js)", 2)
section(doc,
    "The Node.js server runs on port 5000 in development and is deployed to Railway in production. "
    "Its base URL in production is the Railway service URL set in the frontend's config.",
    "Ang Node.js server ay tumatakbo sa port 5000 sa development at naka-deploy sa Railway sa production. "
    "Ang base URL nito sa production ay ang Railway service URL na naka-set sa frontend config."
)

node_endpoints = [
    ("POST /signup", "Creates a new user account and sends a verification email.",
     "Gumagawa ng bagong user account at nagpapadala ng verification email."),
    ("GET  /verify-email?token=...", "Activates a user account after email verification.",
     "Ina-activate ang user account pagkatapos ng email verification."),
    ("POST /login", "Authenticates a user and returns a JWT token.",
     "Nag-a-authenticate ng user at nagbabalik ng JWT token."),
    ("POST /forgot-password", "Sends a 6-digit reset code to the user's email.",
     "Nagpapadala ng 6-digit reset code sa email ng user."),
    ("POST /reset-password", "Resets a user's password using the valid reset code.",
     "Nag-re-reset ng password ng user gamit ang valid reset code."),
    ("GET  /auth/github", "Starts GitHub OAuth login flow.",
     "Nagsisimula ng GitHub OAuth login flow."),
    ("GET  /auth/google", "Starts Google OAuth login flow.",
     "Nagsisimula ng Google OAuth login flow."),
    ("GET  /api/users/:id", "Returns a user's profile data.",
     "Nagbabalik ng profile data ng user."),
    ("POST /api/schemas", "Saves a schema to the database.",
     "Nagse-save ng schema sa database."),
    ("GET  /api/schemas/:userId", "Returns all saved schemas for a user.",
     "Nagbabalik ng lahat ng saved schemas ng user."),
    ("GET  /api/schema/:id", "Returns a single saved schema by its ID.",
     "Nagbabalik ng isang saved schema base sa ID nito."),
    ("GET  /api/admin/stats", "Returns system-wide statistics (admin only).",
     "Nagbabalik ng system-wide statistics (admin lang)."),
    ("GET  /api/admin/analytics", "Returns usage analytics over time (admin only).",
     "Nagbabalik ng usage analytics sa paglipas ng panahon (admin lang)."),
    ("GET  /api/admin/users", "Returns paginated user list (admin only).",
     "Nagbabalik ng paginated na listahan ng users (admin lang)."),
    ("POST /api/llm/generate-schema", "MAIN LLM ENDPOINT — takes user prompt → returns JSON schema definition.",
     "PANGUNAHING LLM ENDPOINT — tumatanggap ng user prompt → nagbabalik ng JSON schema definition."),
    ("POST /api/llm/suggest-field", "Suggests type and constraints for a single field name using Claude.",
     "Nagmumungkahi ng type at constraints para sa isang field name gamit ang Claude."),
    ("POST /api/llm/augment-schema", "Given a real schema + user prompt, identifies missing fields and generates them with AI.",
     "Binibigyan ng real schema + user prompt, kinikilala ang mga nawawalang fields at ginegenerahan sila ng AI."),
    ("POST /api/llm/expand-search-query", "Expands a search query into related terms using Claude Haiku.",
     "Nagpapalawak ng search query sa mga kaugnay na termino gamit ang Claude Haiku."),
    ("POST /api/datasets", "Saves dataset metadata to the database.",
     "Nagse-save ng dataset metadata sa database."),
]

for method_path, eng, tag in node_endpoints:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(method_path)
    r.bold = True; r.font.name = "Courier New"; r.font.color.rgb = RGBColor(0x0f, 0x76, 0x6e)
    p.add_run(f"\n    EN: {eng}")
    p.add_run(f"\n    TL: {tag}")
    p.paragraph_format.space_after = Pt(4)

add_heading(doc, "3.2 Python FastAPI Endpoints (backend/python/main.py)", 2)
section(doc,
    "The Python service runs on port 8000 and is a separate Railway service. "
    "It handles all heavy data work — dataset downloads, schema analysis, and data generation.",
    "Ang Python service ay tumatakbo sa port 8000 at isang hiwalay na Railway service. "
    "Nangangasiwa ito ng lahat ng mabibigat na data work — dataset downloads, schema analysis, at data generation."
)

python_endpoints = [
    ("POST /api/kaggle/search", "Searches Kaggle for datasets matching a query.",
     "Naghahanap ng datasets sa Kaggle na tugma sa query."),
    ("POST /api/kaggle/download", "Downloads a Kaggle dataset and returns its schema.",
     "Nagda-download ng Kaggle dataset at nagbabalik ng schema nito."),
    ("POST /api/huggingface/search", "Searches HuggingFace for datasets.",
     "Naghahanap ng datasets sa HuggingFace."),
    ("POST /api/huggingface/download", "Downloads a HuggingFace dataset.",
     "Nagda-download ng HuggingFace dataset."),
    ("POST /api/uci/search", "Searches the UCI ML Repository.",
     "Naghahanap sa UCI ML Repository."),
    ("POST /api/uci/download", "Downloads a UCI dataset using the ucimlrepo Python package.",
     "Nagda-download ng UCI dataset gamit ang ucimlrepo Python package."),
    ("POST /api/openml/search", "Searches OpenML for datasets.",
     "Naghahanap ng datasets sa OpenML."),
    ("POST /api/openml/download", "Downloads an OpenML dataset.",
     "Nagda-download ng OpenML dataset."),
    ("POST /api/datagov_ph/search", "Searches Data.gov.ph for Philippine government datasets.",
     "Naghahanap ng Philippine government datasets sa Data.gov.ph."),
    ("POST /api/psa/search", "Searches the Philippine Statistics Authority dataset portal.",
     "Naghahanap sa Philippine Statistics Authority dataset portal."),
    ("POST /api/upload-dataset", "Accepts a user-uploaded CSV, analyzes it, and returns its schema.",
     "Tumatanggap ng user-uploaded CSV, nag-a-analyze nito, at nagbabalik ng schema nito."),
    ("POST /api/generate", "CTGAN PATH — trains Gaussian Copula on a downloaded real dataset and generates rows.",
     "CTGAN PATH — nagsasanay ng Gaussian Copula sa downloaded real dataset at gumagawa ng rows."),
    ("GET  /api/preview/:dataset_id", "Returns first N rows of a generated CSV for in-app preview.",
     "Nagbabalik ng unang N rows ng generated CSV para sa in-app preview."),
    ("POST /api/generate-from-schema", "SCHEMA PATH — generates up to 200 rows directly from a schema definition (no real dataset).",
     "SCHEMA PATH — gumagawa ng hanggang 200 rows direkta mula sa schema definition (walang real dataset)."),
    ("POST /api/generate-multi-table", "Generates multiple related tables with FK consistency and returns a ZIP/XLSX.",
     "Gumagawa ng maramihang kaugnay na tables na may FK consistency at nagbabalik ng ZIP/XLSX."),
    ("POST /api/expand-template", "SCALE-UP PATH — trains Gaussian Copula on the 200-row template to produce large datasets.",
     "SCALE-UP PATH — nagsasanay ng Gaussian Copula sa 200-row template para gumawa ng malalaking datasets."),
    ("GET  /api/download-template/:dataset_id", "Downloads the 200-row template in CSV, JSON, or XLSX.",
     "Nagda-download ng 200-row template sa CSV, JSON, o XLSX."),
    ("POST /api/smart-search", "Searches all 6 dataset sources simultaneously using expanded query terms.",
     "Naghahanap sa lahat ng 6 dataset sources nang sabay-sabay gamit ang expanded query terms."),
    ("POST /api/smart-search/preview", "Returns a preview of a specific search result's actual data.",
     "Nagbabalik ng preview ng aktwal na data ng isang search result."),
]

for method_path, eng, tag in python_endpoints:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(method_path)
    r.bold = True; r.font.name = "Courier New"; r.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)
    p.add_run(f"\n    EN: {eng}")
    p.add_run(f"\n    TL: {tag}")
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — LLM INTEGRATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. LLM Integration — How AI (Claude) Is Used", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

section(doc,
    "SynthCS integrates Anthropic's Claude AI model (specifically Claude Haiku 4.5) at four key points in the application. "
    "All LLM calls are made from the Node.js backend (server.js) using the @anthropic-ai/sdk package — "
    "the API key is stored as an environment variable on Railway.",
    "Ang SynthCS ay gumagamit ng Anthropic Claude AI model (specifically Claude Haiku 4.5) sa apat na pangunahing punto "
    "sa application. Lahat ng LLM calls ay ginagawa mula sa Node.js backend (server.js) gamit ang @anthropic-ai/sdk package — "
    "ang API key ay naka-store bilang environment variable sa Railway."
)

add_heading(doc, "4.1 Schema Generation from User Prompt", 2)
section(doc,
    "LOCATION: server.js → POST /api/llm/generate-schema (line ~925)\n"
    "WHAT IT DOES: The user types a plain-English description of their dataset "
    "(e.g. 'employee payroll with salary by role and department'). The Node.js server sends this to Claude Haiku "
    "with a detailed system prompt that instructs the model to return a structured JSON schema — including field names, "
    "types, constraints (min/max, enum values, date ranges), and descriptions.\n"
    "Claude's response is parsed and returned to the frontend, which populates the Schema Editor.",
    "LOKASYON: server.js → POST /api/llm/generate-schema (linya ~925)\n"
    "GINAGAWA NITO: Nag-type ang user ng plain-English na paglalarawan ng kanilang dataset "
    "(hal. 'employee payroll with salary by role and department'). Ipinapadala ng Node.js server ito kay Claude Haiku "
    "na may detalyadong system prompt na nag-uutos sa model na magbalik ng structured JSON schema — kasama ang field names, "
    "types, constraints (min/max, enum values, date ranges), at descriptions.\n"
    "Ang tugon ni Claude ay pina-parse at ibinalik sa frontend, na nagpupuno sa Schema Editor."
)

add_heading(doc, "Key rules embedded in the system prompt:", 3)
rules = [
    ("ENUM EXTRACTION RULE", "If the user's prompt lists specific allowed values (e.g. 'status: Active, Inactive'), Claude must capture them as enum_values.",
     "Kung ang prompt ng user ay naglalaman ng specific na allowed values (hal. 'status: Active, Inactive'), dapat itong i-capture ni Claude bilang enum_values."),
    ("RANGE RULE", "If the user specifies a numeric range (e.g. 'age between 18 and 65'), Claude sets min_val=18, max_val=65.",
     "Kung nagtakda ang user ng numeric range (hal. 'age between 18 and 65'), itinatakda ni Claude ang min_val=18, max_val=65."),
    ("MULTI-ENTITY RULE", "If the prompt mentions multiple entities (users, orders, products), Claude generates fields for ALL of them using prefix naming.",
     "Kung binabanggit ng prompt ang maramihang entities (users, orders, products), gumagawa si Claude ng fields para sa LAHAT gamit ang prefix naming."),
    ("LOCALE RULE", "If a Philippine location is mentioned, phone/address fields are given descriptions that trigger Philippine-format generation.",
     "Kung binabanggit ang Philippine location, binibigyan ang phone/address fields ng descriptions na nag-ti-trigger ng Philippine-format generation."),
    ("MODERATION", "Inappropriate prompts are blocked. Users receive 1 strike per violation; at 3 strikes the account is permanently banned.",
     "Ang mga inappropriate prompts ay hinaharangan. Nakatatanggap ang users ng 1 strike bawat paglabag; sa 3 strikes ay permanenteng bina-ban ang account."),
]
for rule_name, eng, tag in rules:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(f"• {rule_name}: ")
    r.bold = True
    p.add_run(f"EN: {eng}")
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.6)
    p2.add_run(f"TL: {tag}")

add_heading(doc, "4.2 AI Field Suggestion (Sparkle Button)", 2)
section(doc,
    "LOCATION: server.js → POST /api/llm/suggest-field (line ~1074)\n"
    "WHAT IT DOES: The user clicks the sparkle (✦) icon on any field in the Schema Editor. "
    "The field name is sent to Claude, which returns a recommended type, description, and constraints for that specific field. "
    "This saves users from having to configure fields manually.",
    "LOKASYON: server.js → POST /api/llm/suggest-field (linya ~1074)\n"
    "GINAGAWA NITO: Nag-click ang user sa sparkle (✦) icon sa kahit anong field sa Schema Editor. "
    "Ipinapadala ang field name kay Claude, na nagbabalik ng inirerekomendang type, description, at constraints para sa field na iyon. "
    "Nakakatipid ito ng oras ng user sa manual na pag-configure ng fields."
)

add_heading(doc, "4.3 Schema Augmentation (LLM + Real Dataset Hybrid)", 2)
section(doc,
    "LOCATION: server.js → POST /api/llm/augment-schema (line ~1125)\n"
    "WHAT IT DOES: This is the HYBRID path. The system already has a real dataset's schema. "
    "Claude is asked: 'Here is what the real dataset has. Here is what the user's prompt asked for. "
    "What fields are MISSING?' Claude returns only the missing fields as AI-generated additions. "
    "These are tagged with a yellow badge in the Schema Editor so users can tell which fields are real vs. AI.",
    "LOKASYON: server.js → POST /api/llm/augment-schema (linya ~1125)\n"
    "GINAGAWA NITO: Ito ang HYBRID na landas. Mayroon na ang sistema ng schema ng real dataset. "
    "Tinatanong si Claude: 'Narito ang mayroon ang real dataset. Narito ang hinihingi ng prompt ng user. "
    "Anong mga fields ang NAWAWALA?' Nagbabalik si Claude ng mga nawawalang fields bilang AI-generated additions. "
    "Ang mga ito ay may yellow badge sa Schema Editor para malaman ng users kung alin ang real vs. AI."
)

add_heading(doc, "4.4 Search Query Expansion", 2)
section(doc,
    "LOCATION: server.js → POST /api/llm/expand-search-query (line ~1279)\n"
    "WHAT IT DOES: Before searching the 6 dataset sources, the system asks Claude Haiku to expand the user's "
    "search query into 5-6 related terms (synonyms, domain keywords). This increases the chance of finding "
    "relevant datasets even if they are not named exactly what the user typed. "
    "Example: 'student performance' → ['academic achievement', 'grades', 'GPA', 'attendance', 'test scores']",
    "LOKASYON: server.js → POST /api/llm/expand-search-query (linya ~1279)\n"
    "GINAGAWA NITO: Bago maghanap sa 6 dataset sources, tinatanong ng sistema si Claude Haiku na palawakin ang "
    "search query ng user sa 5-6 kaugnay na termino (synonyms, domain keywords). Pinapataas nito ang tsansa na "
    "mahanap ang mga relevant datasets kahit hindi sila pinangalanan ng eksaktong nai-type ng user. "
    "Halimbawa: 'student performance' → ['academic achievement', 'grades', 'GPA', 'attendance', 'test scores']"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — SYNTHETIC DATA ALGORITHM
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Synthetic Data Generation — The Algorithm", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

section(doc,
    "SynthCS uses TWO different generation engines depending on the path the user takes. "
    "Understanding which engine is used when is one of the most important things to know for the panel defense.",
    "Ang SynthCS ay gumagamit ng DALAWANG iba't ibang generation engines depende sa landas na pinili ng user. "
    "Ang pag-unawa kung aling engine ang ginagamit ay isa sa pinakamahalagang bagay na dapat malaman para sa panel defense."
)

add_heading(doc, "5.1 Path A — Schema-Based Generation (No Real Dataset)", 2)
section(doc,
    "FILE: backend/python/smart_gen_data.py (function: gen_col)\n"
    "USED WHEN: The user types a prompt and either no real dataset is found, or the user clicks 'Skip — use pure AI'.",
    "FILE: backend/python/smart_gen_data.py (function: gen_col)\n"
    "GINAGAMIT KUNG: Nag-type ang user ng prompt at wala itong nahanap na real dataset, o nag-click ang user ng 'Skip — use pure AI'."
)

section(doc,
    "HOW IT WORKS (Step by Step):\n"
    "1. The LLM generates a schema JSON with field names, types, and constraints.\n"
    "2. For each column, gen_col() is called with the field type, constraints, and description.\n"
    "3. gen_col() generates realistic values column by column:\n"
    "   - string with enum_values → randomly samples from the allowed list\n"
    "   - integer/float → generates numbers within min/max using the specified distribution (uniform, normal, skewed)\n"
    "   - boolean → generates True/False at the specified true_ratio\n"
    "   - date → random dates between date_from and date_to\n"
    "   - email/phone/name/address → uses Philippine-aware pools (PH names, PH streets, PH mobile prefixes)\n"
    "4. NULL injection: a null_rate % of values are replaced with None.\n"
    "5. The result is a 200-row DataFrame saved as template.csv.",
    "PAANO ITO GUMAGANA (Sunud-sunod):\n"
    "1. Gumagawa ang LLM ng schema JSON na may field names, types, at constraints.\n"
    "2. Para sa bawat column, tinatawag ang gen_col() na may field type, constraints, at description.\n"
    "3. Gumagawa ang gen_col() ng realistic na values column by column:\n"
    "   - string na may enum_values → random sampling mula sa allowed list\n"
    "   - integer/float → gumagawa ng numbers sa loob ng min/max gamit ang tinukoy na distribution\n"
    "   - boolean → gumagawa ng True/False sa tinukoy na true_ratio\n"
    "   - date → random dates sa pagitan ng date_from at date_to\n"
    "   - email/phone/name/address → gumagamit ng Philippine-aware pools (PH names, PH streets, PH mobile prefixes)\n"
    "4. NULL injection: nil_rate % ng values ay pinapalitan ng None.\n"
    "5. Ang resulta ay isang 200-row DataFrame na sine-save bilang template.csv."
)

add_heading(doc, "5.2 Path B — Gaussian Copula Synthesizer (Real Dataset Base)", 2)
section(doc,
    "FILE: backend/python/generator.py (function: _gaussian_copula_sample)\n"
    "USED WHEN: The user selects a real dataset (Kaggle, UCI, etc.) OR when the 200-row template is expanded to a large dataset.",
    "FILE: backend/python/generator.py (function: _gaussian_copula_sample)\n"
    "GINAGAMIT KUNG: Pinili ng user ang isang real dataset (Kaggle, UCI, atbp.) O kapag pinalawak ang 200-row template sa malaking dataset."
)

section(doc,
    "WHY GAUSSIAN COPULA (not a neural network):\n"
    "A Gaussian Copula is a statistical method that captures correlations between columns "
    "without needing thousands of training rows. Neural-based methods like CTGAN require "
    "large training datasets and GPU resources. The Gaussian Copula works well on datasets "
    "as small as 200 rows, making it practical for a thesis-scale system running on free-tier cloud.",
    "BAKIT GAUSSIAN COPULA (hindi neural network):\n"
    "Ang Gaussian Copula ay isang statistical method na kumukuha ng mga kaugnayan (correlations) "
    "sa pagitan ng mga columns nang hindi nangangailangan ng libu-libong training rows. Ang mga "
    "neural-based methods tulad ng CTGAN ay nangangailangan ng malalaking training datasets at GPU resources. "
    "Ang Gaussian Copula ay gumagana nang mabuti sa datasets na kasing liit ng 200 rows, "
    "na ginagawa itong praktikal para sa isang thesis-scale system na tumatakbo sa free-tier cloud."
)

add_heading(doc, "How the Gaussian Copula works (Step by Step):", 3)
steps_eng = [
    "1. CATEGORICAL COLUMNS: Each text/category column's value frequency distribution is learned. "
       "New rows sample from this distribution — if 60% of 'status' values are 'Active', the synthetic data keeps that 60%.",
    "2. NUMERIC COLUMNS — MARGINAL DISTRIBUTION: Each numeric column's distribution is learned via its empirical CDF (Cumulative Distribution Function). "
       "This captures the real shape of the data — skewed distributions, heavy tails, etc.",
    "3. CORRELATION CAPTURE: All numeric columns are converted to standard normal scores (Gaussian space). "
       "A correlation matrix is computed. This captures HOW columns relate — e.g., if salary increases with experience in real data, "
       "that relationship is captured.",
    "4. CHOLESKY DECOMPOSITION: The correlation matrix is decomposed using Cholesky decomposition (L = cholesky(R)). "
       "This is the mathematical heart of the Gaussian Copula. It produces a lower triangular matrix L such that L × L^T = R.",
    "5. CORRELATED SAMPLING: New correlated normal samples Z are generated as Z = random_normals @ L^T. "
       "These are converted back to uniform [0,1] using the normal CDF, then to original value scale "
       "via empirical quantile interpolation.",
    "6. DATE HANDLING: Date columns are converted to ordinal integers before processing and converted back after generation.",
    "7. USER CHANGES: The user's schema edits (renames, type changes, not-null constraints) are applied as a post-processing step.",
]
steps_tag = [
    "1. MGA CATEGORICAL COLUMNS: Natututo ang sistema sa frequency distribution ng bawat text/category column. "
       "Ang mga bagong row ay kumukuha mula sa distribution na ito — kung 60% ng 'status' values ay 'Active', "
       "pinapanatili ng synthetic data ang 60% na iyon.",
    "2. MGA NUMERIC COLUMNS — MARGINAL DISTRIBUTION: Natututo ang sistema sa distribution ng bawat numeric column "
       "sa pamamagitan ng empirical CDF nito. Kumukuha ito ng tunay na hugis ng data — skewed distributions, mabibigat na tails, atbp.",
    "3. CORRELATION CAPTURE: Lahat ng numeric columns ay kino-convert sa standard normal scores (Gaussian space). "
       "Kinukuha ang correlation matrix. Kumukuha ito ng KAUGNAYAN ng mga columns — hal., kung tumataas ang salary sa "
       "tataas na experience sa real data, nakuha ang relasyong iyon.",
    "4. CHOLESKY DECOMPOSITION: Ang correlation matrix ay dina-decompose gamit ang Cholesky decomposition (L = cholesky(R)). "
       "Ito ang mathematical na puso ng Gaussian Copula. Gumagawa ito ng lower triangular matrix L kung saan L × L^T = R.",
    "5. CORRELATED SAMPLING: Ang mga bagong correlated normal samples Z ay ginagawa bilang Z = random_normals @ L^T. "
       "Kino-convert ang mga ito pabalik sa uniform [0,1] gamit ang normal CDF, pagkatapos sa original value scale "
       "sa pamamagitan ng empirical quantile interpolation.",
    "6. DATE HANDLING: Ang mga date columns ay kino-convert sa ordinal integers bago i-process at kino-convert pabalik pagkatapos ng generation.",
    "7. USER CHANGES: Ang mga pag-edit ng user sa schema (renames, type changes, not-null constraints) ay inilalapat bilang post-processing step.",
]
for eng_s, tag_s in zip(steps_eng, steps_tag):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run("EN: ")
    r.bold = True; r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p.add_run(eng_s)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.6)
    r2 = p2.add_run("TL: ")
    r2.bold = True; r2.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p2.add_run(tag_s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — MULTI-TABLE GENERATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Multi-Table Generation", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

section(doc,
    "FILE: backend/python/main.py → POST /api/generate-multi-table (line ~1018)\n"
    "SUPPORTING FILE: backend/python/relational_gen.py",
    "FILE: backend/python/main.py → POST /api/generate-multi-table (linya ~1018)\n"
    "SUPPORTING FILE: backend/python/relational_gen.py"
)

section(doc,
    "HOW IT WORKS:\n"
    "1. TOPOLOGICAL SORT: The system builds a dependency graph from foreign key (FK) relationships. "
       "Tables are sorted so parent tables (the ones being referenced) are always generated BEFORE child tables. "
       "Algorithm used: Kahn's algorithm for topological sorting.\n"
    "2. PARENT TABLE GENERATION: Parent tables (no FK dependencies) are generated first using schema-based generation.\n"
    "3. FK COLUMN FILLING: For child tables, FK columns are filled by randomly sampling from the ALREADY-GENERATED parent table's values. "
       "This ensures referential integrity — an Orders.user_id will always be a valid ID from the Users table.\n"
    "4. HR PAYROLL CONSISTENCY: If a parent table has role/job columns and a child table has salary/experience columns, "
       "the system automatically makes salary consistent with role — a Director always earns more than an Intern.\n"
    "5. OUTPUT: All generated tables are bundled into a ZIP file (one CSV per table) or a single XLSX file.",
    "PAANO ITO GUMAGANA:\n"
    "1. TOPOLOGICAL SORT: Nagtatayo ang sistema ng dependency graph mula sa foreign key (FK) relationships. "
       "Iniayos ang mga tables para laging nagagawa ang parent tables (ang mga na-re-reference) BAGO ang child tables. "
       "Algorithm na ginamit: Kahn's algorithm para sa topological sorting.\n"
    "2. PARENT TABLE GENERATION: Ang mga parent tables (walang FK dependencies) ay ginagawa muna gamit ang schema-based generation.\n"
    "3. FK COLUMN FILLING: Para sa child tables, ang mga FK columns ay pinupuno sa pamamagitan ng random sampling mula sa mga values "
       "ng NAUNANG NAGAWANG parent table. Tinitiyak nito ang referential integrity — ang Orders.user_id ay laging magiging valid na ID mula sa Users table.\n"
    "4. HR PAYROLL CONSISTENCY: Kung ang parent table ay may role/job columns at ang child table ay may salary/experience columns, "
       "awtomatikong ginagawang consistent ng sistema ang salary sa role — ang Director ay laging kumikita nang higit sa Intern.\n"
    "5. OUTPUT: Lahat ng nagawang tables ay pinagsamasama sa isang ZIP file (isang CSV bawat table) o isang XLSX file."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DATA FLOW
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Complete Data Flow — From User Prompt to Download", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

add_heading(doc, "7.1 LLM Path (Pure AI — No Real Dataset)", 2)
flow_llm_eng = [
    "User types a description → Frontend sends to Node.js POST /api/llm/generate-schema",
    "Node.js sends prompt + system rules to Claude Haiku → Claude returns JSON schema",
    "Frontend shows Schema Editor with pre-filled fields (types, constraints, descriptions)",
    "User adjusts schema if needed → clicks Preview",
    "Frontend sends schema to Python POST /api/generate-from-schema → Python runs gen_col() per field",
    "Python returns 200-row template → Frontend shows preview",
    "User downloads template (CSV/JSON/XLSX) OR clicks 'Expand' to scale up",
    "If Expand: Frontend sends to Python POST /api/expand-template → Gaussian Copula trains on 200 rows → generates N rows",
    "Python streams the CSV file back → Browser downloads it",
]
flow_llm_tag = [
    "Nag-type ang user ng paglalarawan → Nagpapadala ang Frontend sa Node.js POST /api/llm/generate-schema",
    "Nagpapadala ang Node.js ng prompt + system rules kay Claude Haiku → Nagbabalik si Claude ng JSON schema",
    "Nagpapakita ang Frontend ng Schema Editor na may pre-filled na fields (types, constraints, descriptions)",
    "Iniaayos ng user ang schema kung kinakailangan → nag-click ng Preview",
    "Nagpapadala ang Frontend ng schema sa Python POST /api/generate-from-schema → Tumatakbo ang Python ng gen_col() bawat field",
    "Nagbabalik ang Python ng 200-row template → Nagpapakita ang Frontend ng preview",
    "Dina-download ng user ang template (CSV/JSON/XLSX) O nag-click ng 'Expand' para palakihin",
    "Kung Expand: Nagpapadala ang Frontend sa Python POST /api/expand-template → Nagsasanay ang Gaussian Copula sa 200 rows → gumagawa ng N rows",
    "Nagst-stream ang Python ng CSV file pabalik → Dina-download ito ng Browser",
]
for i, (e, t) in enumerate(zip(flow_llm_eng, flow_llm_tag)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"Step {i+1} EN: ")
    r.bold = True; r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p.add_run(e)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    r2 = p2.add_run("      TL: ")
    r2.bold = True; r2.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p2.add_run(t)

add_heading(doc, "7.2 Real Dataset Path (Kaggle / UCI / etc.)", 2)
flow_real_eng = [
    "User searches → Frontend → Node.js POST /api/llm/expand-search-query → expanded terms returned",
    "Frontend → Python POST /api/smart-search → searches all 6 sources simultaneously → returns ranked results",
    "User selects a dataset → Frontend → Python download endpoint (e.g. /api/kaggle/download)",
    "Python downloads CSV, saves to temp_datasets/{uuid}/, runs analyze_dataset() → returns schema",
    "Optional: Frontend → Node.js POST /api/llm/augment-schema → Claude adds missing fields",
    "User adjusts schema → clicks Generate",
    "Frontend → Python POST /api/generate → generate_synthetic_data() runs Gaussian Copula on real CSV",
    "Post-processing: temporal rules, relationship rules, anomaly injection applied",
    "Python streams output CSV → Browser downloads",
]
flow_real_tag = [
    "Naghahanap ang user → Frontend → Node.js POST /api/llm/expand-search-query → naibabalik ang expanded terms",
    "Frontend → Python POST /api/smart-search → naghahanap sa lahat ng 6 sources nang sabay-sabay → nagbabalik ng ranked results",
    "Pumili ang user ng dataset → Frontend → Python download endpoint (hal. /api/kaggle/download)",
    "Nagda-download ang Python ng CSV, nagse-save sa temp_datasets/{uuid}/, nagta-run ng analyze_dataset() → nagbabalik ng schema",
    "Opsyonal: Frontend → Node.js POST /api/llm/augment-schema → nagdadagdag si Claude ng mga nawawalang fields",
    "Iniaayos ng user ang schema → nag-click ng Generate",
    "Frontend → Python POST /api/generate → nagta-run ang generate_synthetic_data() ng Gaussian Copula sa real CSV",
    "Post-processing: temporal rules, relationship rules, anomaly injection inilalapat",
    "Nagst-stream ang Python ng output CSV → Dina-download ng Browser",
]
for i, (e, t) in enumerate(zip(flow_real_eng, flow_real_tag)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"Step {i+1} EN: ")
    r.bold = True; r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p.add_run(e)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    r2 = p2.add_run("      TL: ")
    r2.bold = True; r2.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p2.add_run(t)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — PANEL Q&A
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Likely Panel Questions and Answers", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

qas = [
    (
        "Q: What algorithm do you use to generate synthetic data? Is it CTGAN?",
        "Tanong: Anong algorithm ang ginagamit ninyo para gumawa ng synthetic data? CTGAN ba ito?",
        "A: We use a Gaussian Copula synthesizer, not CTGAN. CTGAN is a deep learning model that requires a large training dataset and GPU. "
        "Our system needs to work on free cloud infrastructure and with datasets as small as 200 rows, so we implemented a Gaussian Copula. "
        "It captures inter-column correlations through a correlation matrix and Cholesky decomposition, and preserves marginal distributions "
        "via empirical CDF mapping. The implementation is in backend/python/generator.py, function _gaussian_copula_sample().",
        "Sagot: Gumagamit kami ng Gaussian Copula synthesizer, hindi CTGAN. Ang CTGAN ay isang deep learning model na nangangailangan ng "
        "malaking training dataset at GPU. Ang aming sistema ay kailangan na gumana sa libreng cloud infrastructure at sa datasets na kasing "
        "liit ng 200 rows, kaya nagpatupad kami ng Gaussian Copula. Kumukuha ito ng mga inter-column correlations sa pamamagitan ng correlation "
        "matrix at Cholesky decomposition, at pinapanatili ang marginal distributions sa pamamagitan ng empirical CDF mapping. "
        "Ang implementation ay nasa backend/python/generator.py, function _gaussian_copula_sample()."
    ),
    (
        "Q: How does the LLM know what schema to generate?",
        "Tanong: Paano alam ng LLM kung anong schema ang gagawin?",
        "A: We send Claude Haiku a carefully engineered system prompt that includes strict rules: field type rules, enum extraction rules, "
        "range rules, multi-entity rules, and locale rules. The model must return a strict JSON format. If the model returns markdown or "
        "explanations, we strip them before parsing. The prompt is in server.js around line 969.",
        "Sagot: Nagpapadala kami kay Claude Haiku ng maingat na engineered na system prompt na may mahigpit na mga panuntunan: field type rules, "
        "enum extraction rules, range rules, multi-entity rules, at locale rules. Ang model ay dapat magbalik ng mahigpit na JSON format. "
        "Kung nagbabalik ang model ng markdown o mga paliwanag, tinatanggal namin ang mga ito bago i-parse. "
        "Ang prompt ay nasa server.js sa paligid ng linya 969."
    ),
    (
        "Q: How do you ensure referential integrity in multi-table generation?",
        "Tanong: Paano ninyo tinitiyak ang referential integrity sa multi-table generation?",
        "A: We perform a topological sort of tables using Kahn's algorithm. Parent tables are always generated before child tables. "
        "When a child table has a foreign key column, we populate it by randomly sampling from the already-generated parent table's "
        "primary key values. This guarantees every foreign key value exists in the parent table. The code is in main.py "
        "generate_multi_table() starting at line 1018.",
        "Sagot: Nagsasagawa kami ng topological sort ng mga tables gamit ang Kahn's algorithm. Ang mga parent tables ay laging ginagawa "
        "bago ang mga child tables. Kapag ang child table ay may foreign key column, pinupuno namin ito sa pamamagitan ng random sampling "
        "mula sa mga primary key values ng naunang nagawang parent table. Tinitiyak nito na ang bawat foreign key value ay umiiral "
        "sa parent table. Ang code ay nasa main.py generate_multi_table() simula sa linya 1018."
    ),
    (
        "Q: What is the difference between the two generation paths?",
        "Tanong: Ano ang pagkakaiba ng dalawang generation paths?",
        "A: Path A (Schema-based): No real dataset. The LLM generates a schema. Python generates values column by column using "
        "data pools and statistical distributions. Output: 200 rows, can be expanded. "
        "Path B (Gaussian Copula): A real downloaded dataset is used as the training base. "
        "The Gaussian Copula learns the statistical structure of the real data and generates new rows that are statistically similar. "
        "The real records are never included in the output — only new synthetic rows.",
        "Sagot: Path A (Schema-based): Walang real dataset. Gumagawa ang LLM ng schema. Gumagawa ang Python ng values column by column "
        "gamit ang data pools at statistical distributions. Output: 200 rows, maaaring palawakin. "
        "Path B (Gaussian Copula): Isang real na na-download na dataset ang ginagamit bilang training base. "
        "Natututo ang Gaussian Copula sa statistical structure ng real data at gumagawa ng mga bagong rows na statistically similar. "
        "Hindi kailanman kasama ang mga real records sa output — bagong synthetic rows lang."
    ),
    (
        "Q: How do you prevent the AI from generating harmful content?",
        "Tanong: Paano ninyo pinipigilan ang AI na gumawa ng mapanganib na nilalaman?",
        "A: We have a keyword-based moderation layer in server.js (function isInappropriatePrompt). "
        "If a prompt contains flagged keywords (violence, weapons, explicit content, etc.), it is blocked. "
        "The user receives a warning strike. After 3 strikes, the account is permanently banned. "
        "This logic runs BEFORE the API call to Claude, so inappropriate prompts never reach the LLM.",
        "Sagot: Mayroon kaming keyword-based moderation layer sa server.js (function isInappropriatePrompt). "
        "Kung ang isang prompt ay naglalaman ng mga naka-flag na keywords (karahasan, sandata, explicit na nilalaman, atbp.), "
        "ito ay hinaharangan. Nakatatanggap ang user ng babala strike. Pagkatapos ng 3 strikes, permanenteng bina-ban ang account. "
        "Ang logic na ito ay tumatakbo BAGO ang API call kay Claude, kaya ang mga inappropriate prompts ay hindi kailanman umaabot sa LLM."
    ),
    (
        "Q: Where is the data stored? Do you store user-uploaded datasets?",
        "Tanong: Saan naka-store ang data? Nino-store ba ninyo ang mga user-uploaded datasets?",
        "A: Downloaded and uploaded datasets are temporarily stored in backend/python/temp_datasets/{uuid}/ on the Railway container's "
        "filesystem. They are NOT stored in the database — only the generated output files. "
        "User accounts, saved schemas, and analytics are stored in a PostgreSQL database hosted on Railway.",
        "Sagot: Ang mga na-download at na-upload na datasets ay pansamantalang inilalagay sa backend/python/temp_datasets/{uuid}/ "
        "sa filesystem ng Railway container. HINDI sila inilalagay sa database — ang mga generated output files lang. "
        "Ang mga user accounts, saved schemas, at analytics ay inilalagay sa isang PostgreSQL database na naka-host sa Railway."
    ),
    (
        "Q: Why did you choose FastAPI for the Python backend instead of Flask or Django?",
        "Tanong: Bakit pinili ninyo ang FastAPI para sa Python backend kaysa sa Flask o Django?",
        "A: FastAPI was chosen for three reasons: (1) It has native async support, which is important for file uploads and "
        "I/O-heavy operations. (2) It auto-generates API documentation. (3) It uses Pydantic models for request/response "
        "validation, which catches data errors early. Flask would have required more manual validation code.",
        "Sagot: Pinili ang FastAPI dahil sa tatlong dahilan: (1) Mayroon itong native async support, na mahalaga para sa file uploads at "
        "I/O-heavy operations. (2) Awtomatiko itong gumagawa ng API documentation. (3) Gumagamit ito ng Pydantic models para sa "
        "request/response validation, na nakakahuli ng mga data errors nang maaga. "
        "Ang Flask ay nangangailangan ng mas maraming manual na validation code."
    ),
    (
        "Q: How does the system handle large datasets (millions of rows)?",
        "Tanong: Paano nangangasiwa ang sistema ng malalaking datasets (milyun-milyong rows)?",
        "A: All downloaded/uploaded datasets are capped at 20,000 rows before processing. "
        "The Gaussian Copula can then be trained efficiently on this sample. For generation, the user can request up to 100,000 output rows. "
        "The system does NOT attempt to process the full dataset if it exceeds 20,000 rows — it takes a stratified random sample instead.",
        "Sagot: Lahat ng mga na-download/na-upload na datasets ay limitado sa 20,000 rows bago i-process. "
        "Ang Gaussian Copula ay maaaring sanayin nang mahusay sa sample na ito. Para sa generation, maaaring humiling ang user ng hanggang "
        "100,000 output rows. Hindi sinusubukan ng sistema na i-process ang buong dataset kung ito ay lalampas sa 20,000 rows — "
        "kumukuha ito ng stratified random sample."
    ),
    (
        "Q: What is the role of the 200-row template?",
        "Tanong: Ano ang papel ng 200-row template?",
        "A: The 200-row template is an intermediate artifact. For the LLM path, it serves two purposes: "
        "(1) It gives the user a quick preview they can check before committing to a large generation. "
        "(2) It acts as the training data for the Gaussian Copula when the user wants to scale up. "
        "Even though 200 rows is a small dataset, the Gaussian Copula can learn statistical distributions from it "
        "because it does not need backpropagation or gradient descent.",
        "Sagot: Ang 200-row template ay isang intermediate artifact. Para sa LLM path, dalawa ang layunin nito: "
        "(1) Nagbibigay ito sa user ng mabilis na preview na maaari nilang suriin bago mag-commit sa malaking generation. "
        "(2) Nagsisilbi itong training data para sa Gaussian Copula kapag gusto ng user na palakihin ang output. "
        "Kahit maliit ang 200 rows, kaya ng Gaussian Copula na matuto ng statistical distributions mula dito "
        "dahil hindi ito nangangailangan ng backpropagation o gradient descent."
    ),
    (
        "Q: How is authentication implemented?",
        "Tanong: Paano ipinatupad ang authentication?",
        "A: We use JWT (JSON Web Tokens) for session management. On login, the server generates a signed JWT using a secret key "
        "stored in environment variables. The frontend stores this token in localStorage. Protected endpoints verify the token "
        "on every request. We also support OAuth via GitHub and Google using Passport.js. "
        "Passwords are hashed with bcrypt (salt rounds: 10) before being stored in PostgreSQL.",
        "Sagot: Gumagamit kami ng JWT (JSON Web Tokens) para sa session management. Sa login, gumagawa ang server ng signed JWT "
        "gamit ang secret key na naka-store sa environment variables. Inilalagay ng frontend ang token na ito sa localStorage. "
        "Bine-verify ng mga protected endpoints ang token sa bawat request. Sumusuporta rin kami ng OAuth sa pamamagitan ng GitHub "
        "at Google gamit ang Passport.js. Ang mga password ay hina-hash gamit ang bcrypt (salt rounds: 10) bago ilalagay sa PostgreSQL."
    ),
]

for q_eng, q_tag, a_eng, a_tag in qas:
    p = doc.add_paragraph()
    r = p.add_run(q_eng)
    r.bold = True; r.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(q_tag)
    r2.bold = True; r2.italic = True; r2.font.color.rgb = RGBColor(0x14, 0x53, 0x2d)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.3)
    r3 = p3.add_run("EN: ")
    r3.bold = True; r3.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p3.add_run(a_eng)

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Inches(0.3)
    r4 = p4.add_run("TL: ")
    r4.bold = True; r4.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p4.add_run(a_tag)

    add_divider(doc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Technology Stack Summary", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

tech = [
    ("React + TypeScript", "Frontend UI framework", "Frontend UI framework"),
    ("Tailwind CSS", "Utility-first CSS styling", "Utility-first CSS styling"),
    ("Vercel", "Frontend deployment (CDN + edge)", "Frontend deployment (CDN + edge)"),
    ("Node.js + Express", "Auth, user management, LLM proxy", "Auth, user management, LLM proxy"),
    ("Python 3.11 + FastAPI", "Data processing, synthesis, dataset search", "Data processing, synthesis, dataset search"),
    ("Railway", "Cloud deployment for both Node.js and Python services", "Cloud deployment para sa Node.js at Python services"),
    ("PostgreSQL", "User accounts, saved schemas, analytics", "User accounts, saved schemas, analytics"),
    ("Anthropic Claude Haiku 4.5", "LLM for schema generation, augmentation, field suggestion, query expansion",
     "LLM para sa schema generation, augmentation, field suggestion, query expansion"),
    ("Gaussian Copula (custom)", "Statistical synthetic data synthesis (backend/python/generator.py)",
     "Statistical synthetic data synthesis (backend/python/generator.py)"),
    ("pandas + numpy + scipy", "Data manipulation, statistical computation", "Data manipulation, statistical computation"),
    ("JWT + bcrypt + Passport.js", "Authentication and session management", "Authentication and session management"),
    ("Kaggle API + ucimlrepo + HuggingFace datasets", "External dataset access", "External dataset access"),
    ("python-multipart", "File upload handling in FastAPI", "File upload handling sa FastAPI"),
    ("openpyxl", "XLSX file generation", "XLSX file generation"),
]

for name, eng_desc, tag_desc in tech:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}: ")
    r.bold = True
    p.add_run(f"EN: {eng_desc}  |  TL: {tag_desc}")
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — LLM + CTGAN INTEGRATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. LLM + CTGAN Integration — How They Work Together", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

section(doc,
    "One of the core contributions of SynthCS is the COMBINATION of two technologies that are normally used separately: "
    "a Large Language Model (Claude) and a statistical synthesizer (Gaussian Copula, referred to as CTGAN in the system). "
    "Neither technology alone is sufficient — the LLM understands domain context but cannot produce statistically "
    "consistent large datasets, while the Gaussian Copula scales data faithfully but has no understanding of what the "
    "data is supposed to represent. Together, they cover each other's weaknesses.",
    "Isa sa mga pangunahing kontribusyon ng SynthCS ay ang PAGSASAMA ng dalawang teknolohiya na karaniwang ginagamit nang hiwalay: "
    "isang Large Language Model (Claude) at isang statistical synthesizer (Gaussian Copula, tinutukoy bilang CTGAN sa sistema). "
    "Walang teknolohiya na sapat nang mag-isa — nauunawaan ng LLM ang domain context ngunit hindi ito makakagawa ng "
    "statistically consistent na malalaking datasets, samantalang sinisigurado ng Gaussian Copula ang statistical fidelity "
    "ngunit walang pag-unawa sa kung ano ang dapat na katawanin ng data. Sama-sama, tinatakpan nila ang kahinaan ng isa't isa."
)

add_heading(doc, "10.1 The Problem Each Technology Solves", 2)

bullet(doc,
    "WHAT THE LLM SOLVES — Domain Understanding: When a user says 'generate a dataset for a gift shop in Olongapo,' "
    "the LLM knows that the dataset should have fields like product_name, price (in Philippine pesos), barangay, "
    "occasion_tags, and vendor_name. It knows realistic value ranges (price: ₱150–₱3,500), enumerations "
    "(occasions: Birthday, Christmas, Fiesta), and locale context (Philippine addresses, Filipino names). "
    "No statistical algorithm can derive this from thin air.",
    "ANO ANG NIRERESOLBA NG LLM — Domain Understanding: Kapag sinabi ng user na 'gumawa ng dataset para sa gift shop sa Olongapo,' "
    "alam ng LLM na ang dataset ay dapat may mga fields tulad ng product_name, price (sa Philippine pesos), barangay, "
    "occasion_tags, at vendor_name. Alam nito ang realistic na value ranges (price: ₱150–₱3,500), enumerations "
    "(okasyon: Birthday, Christmas, Fiesta), at locale context (Philippine addresses, Filipino names). "
    "Walang statistical algorithm ang makakalikha nito mula sa wala."
)

bullet(doc,
    "WHAT THE GAUSSIAN COPULA SOLVES — Statistical Fidelity at Scale: The LLM can only generate a small template "
    "(200 rows). If the user needs 10,000 rows, simply repeating LLM calls would be slow, expensive, and would produce "
    "rows that do not respect correlations between columns (e.g., higher-priced products not necessarily having higher ratings). "
    "The Gaussian Copula learns the statistical structure of the 200-row template — column distributions and "
    "inter-column correlations — and generates thousands of new rows that match that structure exactly.",
    "ANO ANG NIRERESOLBA NG GAUSSIAN COPULA — Statistical Fidelity sa Malaking Bilang: Kaya lamang ng LLM na gumawa "
    "ng maliit na template (200 rows). Kung kailangan ng user ng 10,000 rows, ang paulit-ulit na pagtawag ng LLM ay "
    "magiging mabagal, mahal, at magbubunga ng mga rows na hindi gumagalang sa mga correlations sa pagitan ng mga columns "
    "(hal., mas mahal na produkto ay hindi nangangailangan ng mas mataas na rating). Natututo ang Gaussian Copula sa "
    "statistical structure ng 200-row template — column distributions at inter-column correlations — at gumagawa ng libu-libong "
    "bagong rows na eksaktong tugma sa istrukturang iyon."
)

add_heading(doc, "10.2 The Integrated Pipeline (Step by Step)", 2)

section(doc,
    "This is the exact sequence of events when a user takes the LLM path in SynthCS:",
    "Ito ang eksaktong pagkakasunud-sunod ng mga pangyayari kapag ginagamit ng user ang LLM path sa SynthCS:"
)

pipeline_steps_eng = [
    ("STEP 1 — USER PROMPT",
     "The user types a plain-English description of the dataset they need. "
     "Example: 'Gift shop sales data for Olongapo City with products, customers, and orders.'"),
    ("STEP 2 — LLM SCHEMA GENERATION (Node.js → Claude Haiku)",
     "The Node.js backend sends the prompt to Claude Haiku with a structured system prompt. "
     "Claude returns a JSON schema: field names, types, constraints (min/max, enum values, date ranges), "
     "and descriptions. This JSON defines WHAT the data should look like — not the data itself. "
     "CODE: server.js, POST /api/llm/generate-schema"),
    ("STEP 3 — SCHEMA EDITOR REVIEW",
     "The frontend displays the schema in the Schema Editor. The user can rename fields, "
     "change types, add or remove columns, and adjust constraints. The LLM-generated fields "
     "are highlighted in purple so the user knows they were AI-generated."),
    ("STEP 4 — 200-ROW TEMPLATE GENERATION (Python → smart_gen_data.py)",
     "When the user clicks Generate, the frontend sends the schema to the Python backend "
     "POST /api/generate-from-schema. Python calls gen_col() for each field, generating "
     "200 realistic rows using data pools (Philippine names, addresses, mobile prefixes), "
     "statistical distributions, and the constraints Claude specified. "
     "Output: a 200-row CSV saved as template.csv. "
     "This step is purely rule-based — NO AI is involved here."),
    ("STEP 5 — GAUSSIAN COPULA TRAINING (Python → generator.py)",
     "When the user requests a large dataset, the frontend calls Python POST /api/expand-with-ctgan. "
     "The Gaussian Copula reads the 200-row template.csv as its TRAINING DATA. "
     "It learns: (a) the distribution of each numeric column via empirical CDF, "
     "(b) the frequency distribution of each categorical column, "
     "(c) the correlations between all numeric columns via a correlation matrix. "
     "This training step runs entirely in Python using numpy and scipy — no GPU, no deep learning."),
    ("STEP 6 — CORRELATED SYNTHETIC DATA GENERATION",
     "Using the learned statistical structure, the Gaussian Copula generates N rows "
     "(up to 100,000) of new synthetic data. "
     "The key property: the synthetic rows are STATISTICALLY SIMILAR to the template but are NOT copies. "
     "Column correlations are preserved — if the template shows that higher-priced orders "
     "tend to have higher ratings, the 10,000-row output will maintain that relationship. "
     "CODE: generator.py, function _gaussian_copula_sample()"),
    ("STEP 7 — DATA PREVIEW AND EXPORT",
     "The generated dataset is saved to temp_datasets/{dataset_id}/synthetic_output.csv. "
     "The frontend navigates to the DataPreview page where the user can inspect the data, "
     "view validation metrics (Wasserstein distance, correlation score, ML utility, privacy score), "
     "and download in CSV, JSON, XLSX, SQL, or JSONL format."),
]

pipeline_steps_tag = [
    ("HAKBANG 1 — PROMPT NG USER",
     "Nag-type ang user ng plain-English na paglalarawan ng dataset na kailangan nila. "
     "Halimbawa: 'Gift shop sales data para sa Olongapo City na may mga produkto, customer, at orders.'"),
    ("HAKBANG 2 — LLM SCHEMA GENERATION (Node.js → Claude Haiku)",
     "Nagpapadala ang Node.js backend ng prompt kay Claude Haiku na may structured system prompt. "
     "Nagbabalik si Claude ng JSON schema: field names, types, constraints (min/max, enum values, date ranges), "
     "at mga paglalarawan. Tinutukoy ng JSON na ito kung ANO ang dapat magmukhang data — hindi ang data mismo. "
     "CODE: server.js, POST /api/llm/generate-schema"),
    ("HAKBANG 3 — SCHEMA EDITOR REVIEW",
     "Ipinapakita ng frontend ang schema sa Schema Editor. Maaaring palitan ng user ang pangalan ng mga fields, "
     "palitan ang mga types, magdagdag o magtanggal ng mga columns, at ayusin ang mga constraints. "
     "Ang mga LLM-generated na fields ay naka-highlight sa purple para malaman ng user na AI ang gumawa ng mga ito."),
    ("HAKBANG 4 — 200-ROW TEMPLATE GENERATION (Python → smart_gen_data.py)",
     "Kapag nag-click ang user ng Generate, nagpapadala ang frontend ng schema sa Python backend "
     "POST /api/generate-from-schema. Tinatawag ng Python ang gen_col() para sa bawat field, "
     "gumagawa ng 200 realistic na rows gamit ang mga data pools (Filipino names, addresses, mobile prefixes), "
     "statistical distributions, at ang mga constraints na tinukoy ni Claude. "
     "Output: isang 200-row CSV na sine-save bilang template.csv. "
     "Ang hakbang na ito ay purong rule-based — WALANG AI na kasangkot dito."),
    ("HAKBANG 5 — GAUSSIAN COPULA TRAINING (Python → generator.py)",
     "Kapag humiling ang user ng malaking dataset, tinatawag ng frontend ang Python POST /api/expand-with-ctgan. "
     "Binabasa ng Gaussian Copula ang 200-row template.csv bilang TRAINING DATA nito. "
     "Natututo ito ng: (a) distribution ng bawat numeric column sa pamamagitan ng empirical CDF, "
     "(b) frequency distribution ng bawat categorical column, "
     "(c) mga correlations sa pagitan ng lahat ng numeric columns sa pamamagitan ng correlation matrix. "
     "Ang training step na ito ay tumatakbo nang buo sa Python gamit ang numpy at scipy — walang GPU, walang deep learning."),
    ("HAKBANG 6 — CORRELATED SYNTHETIC DATA GENERATION",
     "Gamit ang natutunan na statistical structure, gumagawa ang Gaussian Copula ng N rows "
     "(hanggang 100,000) ng bagong synthetic data. "
     "Ang pangunahing katangian: ang mga synthetic rows ay STATISTICALLY SIMILAR sa template ngunit hindi kopya ang mga ito. "
     "Ang mga column correlations ay nananatili — kung ipinakita ng template na ang mas mahal na orders "
     "ay may mas mataas na ratings, mananatili ang relasyong iyon sa 10,000-row output. "
     "CODE: generator.py, function _gaussian_copula_sample()"),
    ("HAKBANG 7 — DATA PREVIEW AT EXPORT",
     "Ang nagawang dataset ay sine-save sa temp_datasets/{dataset_id}/synthetic_output.csv. "
     "Nagna-navigate ang frontend sa DataPreview page kung saan maaaring suriin ng user ang data, "
     "tingnan ang mga validation metrics (Wasserstein distance, correlation score, ML utility, privacy score), "
     "at i-download sa CSV, JSON, XLSX, SQL, o JSONL format."),
]

for (step_eng, detail_eng), (step_tag, detail_tag) in zip(pipeline_steps_eng, pipeline_steps_tag):
    p = doc.add_paragraph()
    r = p.add_run(step_eng)
    r.bold = True; r.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    r2 = p2.add_run("EN: ")
    r2.bold = True; r2.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p2.add_run(detail_eng)
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.3)
    r3 = p3.add_run("TL: ")
    r3.bold = True; r3.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p3.add_run(detail_tag)
    doc.add_paragraph()

add_heading(doc, "10.3 Why This Combination Is Effective", 2)

section(doc,
    "The LLM + Gaussian Copula combination addresses a core challenge in synthetic data generation: "
    "how do you generate data that is both MEANINGFUL (looks like real data from a specific domain) "
    "and STATISTICALLY VALID (preserves distributions and correlations at scale)?\n\n"
    "LLM alone: Can generate domain-aware data but is slow, expensive, and loses statistical "
    "coherence beyond a few hundred rows.\n"
    "Gaussian Copula alone: Can scale efficiently and preserve statistics but needs real training "
    "data and has no domain awareness — it cannot invent a gift shop schema from nothing.\n"
    "COMBINED: The LLM provides the domain-aware seed (200 realistic rows that look exactly like "
    "the target domain). The Gaussian Copula learns from this seed and generates thousands of "
    "statistically consistent rows. Each technology contributes what the other cannot.",
    "Ang kombinasyon ng LLM + Gaussian Copula ay tumutugon sa isang pangunahing hamon sa synthetic data generation: "
    "paano mo gagawing data na parehong MAKABULUHAN (mukhang tunay na data mula sa isang partikular na domain) "
    "at STATISTICALLY VALID (pinapanatili ang mga distributions at correlations sa malaking bilang)?\n\n"
    "LLM lamang: Kayang gumawa ng domain-aware na data ngunit mabagal, mahal, at nawawala ang statistical "
    "coherence nang higit sa ilang daang rows.\n"
    "Gaussian Copula lamang: Kayang mag-scale nang mahusay at mapanatili ang mga statistics ngunit "
    "nangangailangan ng tunay na training data at walang domain awareness — hindi ito makakaimbento ng "
    "gift shop schema mula sa wala.\n"
    "PINAGSAMA: Nagbibigay ang LLM ng domain-aware na binhi (200 realistic na rows na eksaktong "
    "mukhang target na domain). Natututo ang Gaussian Copula mula sa binhini at gumagawa ng libu-libong "
    "statistically consistent na rows. Bawat teknolohiya ay nag-aambag ng hindi kaya ng isa."
)

add_heading(doc, "10.4 Panel Q&A — Combined Integration", 2)

combined_qas = [
    (
        "Q: Why do you call it 'CTGAN' in the system if you use Gaussian Copula?",
        "Tanong: Bakit tinatawag itong 'CTGAN' sa sistema kung Gaussian Copula ang ginagamit ninyo?",
        "A: CTGAN (Conditional Tabular GAN) is the industry-standard term for neural-network-based tabular data synthesis. "
        "We label the expansion endpoint 'expand-with-ctgan' because it plays the same ROLE in the pipeline — "
        "taking a small seed dataset and producing a large synthetic one. However, our actual implementation uses "
        "a Gaussian Copula instead of a GAN because: (1) GANs require GPUs and thousands of training rows, "
        "(2) our seed is only 200 rows, and (3) we run on free-tier cloud. The Gaussian Copula achieves "
        "the same goal — statistical fidelity at scale — without those requirements.",
        "Sagot: Ang CTGAN (Conditional Tabular GAN) ay ang industry-standard na termino para sa neural-network-based "
        "tabular data synthesis. Tinutukoy namin ang expansion endpoint na 'expand-with-ctgan' dahil ginagampanan "
        "nito ang parehong PAPEL sa pipeline — kumukuha ng maliit na seed dataset at gumagawa ng malaking synthetic na isa. "
        "Gayunpaman, ang aming aktwal na implementation ay gumagamit ng Gaussian Copula sa halip na GAN dahil: "
        "(1) Ang mga GAN ay nangangailangan ng GPUs at libu-libong training rows, "
        "(2) ang aming seed ay 200 rows lamang, at (3) tumatakbo kami sa free-tier cloud. "
        "Naabot ng Gaussian Copula ang parehong layunin — statistical fidelity sa malaking bilang — nang walang mga kinakailangang iyon."
    ),
    (
        "Q: What is the contribution of the LLM if the actual data generation is statistical?",
        "Tanong: Ano ang kontribusyon ng LLM kung statistical ang aktwal na data generation?",
        "A: The LLM's contribution is DOMAIN BOOTSTRAPPING. Without the LLM, the Gaussian Copula has nothing to train on "
        "for custom schemas — it needs a starting dataset. The LLM creates that starting point: a 200-row template that is "
        "already domain-correct, field-complete, and constraint-aware. This is the key innovation: instead of needing a "
        "real dataset to train the synthesizer, we use AI to SIMULATE one, then train the synthesizer on that simulation. "
        "The LLM also provides the schema definition that determines what fields exist, what types they are, and what values are valid.",
        "Sagot: Ang kontribusyon ng LLM ay DOMAIN BOOTSTRAPPING. Kung wala ang LLM, walang masasanayan ang Gaussian Copula "
        "para sa custom schemas — kailangan nito ng starting dataset. Ginagawa ng LLM ang starting point na iyon: isang "
        "200-row template na domain-correct na, kumpleto ang fields, at alam ang mga constraints. "
        "Ito ang pangunahing inobasyon: sa halip na kailangan ng tunay na dataset para sanayin ang synthesizer, "
        "ginagamit namin ang AI para GAYAHIN ang isa, pagkatapos sanayin ang synthesizer sa simulation na iyon. "
        "Nagbibigay din ang LLM ng schema definition na nagtatakda kung anong fields ang umiiral, "
        "anong types ang mga ito, at anong values ang valid."
    ),
    (
        "Q: Is the 200-row template generated by the LLM or by a different algorithm?",
        "Tanong: Ang LLM ba ang gumagawa ng 200-row template o ibang algorithm?",
        "A: The LLM generates the SCHEMA (field definitions), NOT the rows. The 200 rows are generated by "
        "backend/python/smart_gen_data.py using a rule-based column generator (gen_col()). "
        "gen_col() uses the schema constraints that the LLM defined — min/max ranges, enum lists, date ranges, "
        "null rates — to produce values for each column. The LLM defines WHAT the data should look like; "
        "smart_gen_data.py does the actual row-by-row generation.",
        "Sagot: Ang LLM ay gumagawa ng SCHEMA (mga kahulugan ng field), HINDI ang mga rows. Ang 200 rows ay ginagawa ng "
        "backend/python/smart_gen_data.py gamit ang isang rule-based column generator (gen_col()). "
        "Ginagamit ng gen_col() ang schema constraints na tinukoy ng LLM — min/max ranges, enum lists, date ranges, "
        "null rates — para makagawa ng values para sa bawat column. Tinutukoy ng LLM kung ANO ang dapat magmukhang data; "
        "ginagawa ng smart_gen_data.py ang aktwal na row-by-row generation."
    ),
    (
        "Q: What happens to the original 200-row template after the Gaussian Copula runs?",
        "Tanong: Ano ang nangyayari sa orihinal na 200-row template pagkatapos tumakbo ang Gaussian Copula?",
        "A: The 200-row template (template.csv) remains on the server as a training artifact. "
        "The Gaussian Copula reads it, learns its statistical structure, then generates a COMPLETELY NEW set of rows "
        "saved as synthetic_output.csv. The template rows themselves are NOT included in the final output — "
        "only the newly generated synthetic rows are returned to the user. "
        "This is important: the final dataset contains zero real/template rows.",
        "Sagot: Ang 200-row template (template.csv) ay nananatili sa server bilang training artifact. "
        "Binabasa ito ng Gaussian Copula, natututo sa statistical structure nito, pagkatapos ay gumagawa ng "
        "GANAP NA BAGONG hanay ng mga rows na sine-save bilang synthetic_output.csv. "
        "Ang mga template rows mismo ay HINDI kasama sa final output — "
        "ang mga bagong nagawang synthetic rows lamang ang ibinalik sa user. "
        "Mahalaga ito: ang final dataset ay naglalaman ng zero real/template rows."
    ),
]

for q_eng, q_tag, a_eng, a_tag in combined_qas:
    p = doc.add_paragraph()
    r = p.add_run(q_eng)
    r.bold = True; r.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(q_tag)
    r2.bold = True; r2.italic = True; r2.font.color.rgb = RGBColor(0x14, 0x53, 0x2d)
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.3)
    r3 = p3.add_run("EN: ")
    r3.bold = True; r3.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    p3.add_run(a_eng)
    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Inches(0.3)
    r4 = p4.add_run("TL: ")
    r4.bold = True; r4.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p4.add_run(a_tag)
    add_divider(doc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — DEMO DATASET PROMPTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "11. Demo Dataset Prompts — For System Defense Demonstration", 1, (0x6d, 0x28, 0xd9))
add_divider(doc)

section(doc,
    "This section provides ready-to-use prompts that demonstrate SynthCS capabilities during the panel defense. "
    "These prompts were chosen because they showcase key features: schema awareness, approval logic, numeric correlations, "
    "and categorical value control. Use these exact prompts when demonstrating the system live.",
    "Ang seksyon na ito ay nagbibigay ng mga handa nang gamitin na prompts na nagpapakita ng mga kakayahan ng SynthCS "
    "sa panahon ng panel defense. Ang mga prompt na ito ay pinili dahil ipinapakita nila ang mga pangunahing features: "
    "schema awareness, approval logic, numeric correlations, at categorical value control. "
    "Gamitin ang mga eksaktong prompt na ito kapag nagde-demonstrate ng sistema nang live."
)

# ─── Demo Prompt 1: Loan Applicant Dataset ────────────────────────────────────
add_heading(doc, "11.1 Demo Prompt — Loan Applicant Dataset", 2)

section(doc,
    "USE CASE: Demonstrates SynthCS generating financial risk data with realistic approval logic. "
    "Suitable for fraud detection, credit scoring, or machine learning classification demos.",
    "GAMIT: Ipinapakita ang SynthCS na gumagawa ng financial risk data na may realistic na approval logic. "
    "Angkop para sa fraud detection, credit scoring, o machine learning classification demos."
)

add_heading(doc, "Prompt to type into SynthCS:", 3)
add_code(doc,
    "Loan applicant dataset with fields: id, loan_amount, term_months, credit_score, "
    "annual_income, employment_years, debt_to_income, home_ownership, purpose, approved"
)

add_heading(doc, "What to say to the panel while typing this prompt:", 3)
section(doc,
    "EN: \"I will now type a plain-English description of the dataset I need. Notice that I am not writing code or SQL — "
    "just a natural description. SynthCS will send this to Claude, which will extract the field names, infer appropriate "
    "types and constraints for each field, and return a structured schema.\"",
    "TL: \"Magta-type na ako ng plain-English na paglalarawan ng dataset na kailangan ko. Pansinin na hindi ako "
    "sumusulat ng code o SQL — natural na paglalarawan lamang. Ipapadala ito ng SynthCS kay Claude, na kukuha ng "
    "mga field names, mag-iinfer ng angkop na types at constraints para sa bawat field, at magbabalik ng structured schema.\""
)

add_heading(doc, "Expected schema fields and their types:", 3)
loan_fields = [
    ("id", "integer", "Unique applicant identifier. Auto-increments from 1.",
     "Natatanging identifier ng aplikante. Awtomatikong dumarami mula sa 1."),
    ("loan_amount", "float", "Requested loan amount in Philippine pesos. Range: ₱50,000–₱5,000,000.",
     "Halagang hinihingi bilang pautang sa Philippine pesos. Range: ₱50,000–₱5,000,000."),
    ("term_months", "integer", "Loan repayment period in months. Enum: 12, 24, 36, 48, 60.",
     "Panahon ng pagbabayad ng pautang sa buwan. Enum: 12, 24, 36, 48, 60."),
    ("credit_score", "integer", "Applicant's credit score. Range: 300–850. Higher = better creditworthiness.",
     "Credit score ng aplikante. Range: 300–850. Mas mataas = mas mahusay na creditworthiness."),
    ("annual_income", "float", "Annual income in Philippine pesos. Range: ₱180,000–₱5,000,000.",
     "Taunang kita sa Philippine pesos. Range: ₱180,000–₱5,000,000."),
    ("employment_years", "float", "Years at current employer. Range: 0–30.",
     "Taon sa kasalukuyang employer. Range: 0–30."),
    ("debt_to_income", "float", "Total monthly debt divided by monthly income. Range: 0.05–0.80.",
     "Kabuuang buwanang utang na hinati sa buwanang kita. Range: 0.05–0.80."),
    ("home_ownership", "string (enum)", "Housing status. Values: OWN, RENT, MORTGAGE.",
     "Katayuan sa pabahay. Values: OWN, RENT, MORTGAGE."),
    ("purpose", "string (enum)", "Reason for the loan. Values: home_improvement, debt_consolidation, education, medical, business, car, vacation.",
     "Dahilan ng pautang. Values: home_improvement, debt_consolidation, education, medical, business, car, vacation."),
    ("approved", "boolean", "Loan approval decision. true = approved, false = declined.",
     "Desisyon sa pag-apruba ng pautang. true = approved, false = declined."),
]

for fname, ftype, eng_desc, tag_desc in loan_fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{fname}  ")
    r.bold = True; r.font.name = "Courier New"; r.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    r2 = p.add_run(f"[{ftype}]")
    r2.font.color.rgb = RGBColor(0x0f, 0x76, 0x6e); r2.font.name = "Courier New"
    p.add_run(f"\n    EN: {eng_desc}\n    TL: {tag_desc}")
    p.paragraph_format.space_after = Pt(4)

add_heading(doc, "Generation settings to use during the demo:", 3)
demo_settings = [
    ("Row count", "50 rows", "Small enough to load instantly, large enough to show variety.",
     "Sapat na maliit para mag-load agad, sapat na malaki para ipakita ang pagkakaiba-iba."),
    ("Generation path", "LLM Path (no real dataset)", "Shows the pure AI schema generation capability.",
     "Ipinapakita ang purong AI schema generation capability."),
    ("Approval logic to mention", "~60% approved, ~40% declined",
     "The Gaussian Copula will learn the true/false ratio from the 200-row template and preserve it at scale.",
     "Matututo ang Gaussian Copula sa true/false ratio mula sa 200-row template at pananatilihin ito sa malaking bilang."),
]
for setting, value, eng_note, tag_note in demo_settings:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"• {setting}: ")
    r.bold = True
    r2 = p.add_run(value)
    r2.bold = True; r2.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    p.add_run(f"\n    EN: {eng_note}\n    TL: {tag_note}")
    p.paragraph_format.space_after = Pt(4)

add_heading(doc, "Key talking point for the panel:", 3)
section(doc,
    "EN: \"The approved column demonstrates how SynthCS can encode business logic into schema constraints. "
    "By setting the true_ratio in the schema, I can control the approval rate in the generated dataset. "
    "The Gaussian Copula then learns the correlation between credit_score, debt_to_income, and approved — "
    "so in the scaled-up output, high-credit-score applicants are statistically more likely to be approved, "
    "just as they would be in a real bank's dataset.\"",
    "TL: \"Ang approved column ay nagpapakita kung paano makakapaglagay ang SynthCS ng business logic sa schema constraints. "
    "Sa pamamagitan ng pagtatakda ng true_ratio sa schema, makokontrol ko ang approval rate sa nagawang dataset. "
    "Pagkatapos ay matututo ang Gaussian Copula sa correlation sa pagitan ng credit_score, debt_to_income, at approved — "
    "kaya sa scaled-up na output, ang mga aplikanteng may mataas na credit score ay statistically mas malamang na maaprubahan, "
    "tulad ng sa tunay na dataset ng isang bangko.\""
)

add_divider(doc)

add_heading(doc, "11.2 How the Generated Dataset Can Be Used (Downstream Application)", 2)
section(doc,
    "The 50-row (or scaled-up) loan applicant dataset generated by SynthCS can be directly used in downstream applications "
    "such as a fraud detection or credit scoring prototype. The following fields are particularly useful for ML classification:",
    "Ang 50-row (o scaled-up) loan applicant dataset na ginawa ng SynthCS ay maaaring direktang gamitin sa mga downstream applications "
    "tulad ng isang fraud detection o credit scoring prototype. Ang mga sumusunod na fields ay partikular na kapaki-pakinabang "
    "para sa ML classification:"
)

ml_fields = [
    ("Features (X)", "credit_score, annual_income, employment_years, debt_to_income, loan_amount, home_ownership, purpose",
     "Mga input na ginagamit ng modelo para hulaan"),
    ("Target (y)", "approved",
     "Ang label na sinusubukan ng modelo na hulaan (binary: true/false)"),
]
for role, fields_list, tag_desc in ml_fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"• {role}: ")
    r.bold = True; r.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    p.add_run(f"{fields_list}\n    TL: {tag_desc}")
    p.paragraph_format.space_after = Pt(4)

section(doc,
    "EN: This demonstrates the VALUE CHAIN of SynthCS: "
    "(1) No real loan data is needed — SynthCS generates realistic synthetic records. "
    "(2) The generated dataset can be downloaded as CSV and immediately imported into a Jupyter notebook, "
    "scikit-learn pipeline, or Tableau dashboard. "
    "(3) Privacy is preserved — no real applicant PII is ever used or exposed.",
    "TL: Ipinapakita nito ang VALUE CHAIN ng SynthCS: "
    "(1) Hindi kailangan ng tunay na loan data — gumagawa ang SynthCS ng realistic synthetic records. "
    "(2) Ang nagawang dataset ay maaaring i-download bilang CSV at agad na i-import sa Jupyter notebook, "
    "scikit-learn pipeline, o Tableau dashboard. "
    "(3) Napanatili ang privacy — walang tunay na PII ng aplikante ang ginamit o na-expose."
)

add_divider(doc)

# ── Section 12: Import & Export Format Design Decision ───────────────────────
add_heading(doc, "12. Import & Export Format Design Decision", 1)

add_heading(doc, "12.1 Why SynthCS Only Accepts CSV for Import", 2)
section(doc,
    "EN: SynthCS supports only CSV file upload for dataset import. "
    "This is an intentional design decision, not a technical limitation. "
    "The synthesis models — CTGAN and Gaussian Copula — both require flat, tabular data: "
    "a simple rows-and-columns structure. "
    "Formats such as JSON or SQL can be hierarchical or multi-table, which would require "
    "additional preprocessing and schema-flattening logic before synthesis could begin. "
    "CSV is also the universal export format across spreadsheets (Excel, Google Sheets), "
    "databases (MySQL Workbench, pgAdmin), and data science tools (pandas, R). "
    "This means users can convert from virtually any source in one step. "
    "The decision prioritized getting the synthesis pipeline right over supporting multiple import formats.",
    "TL: Ang SynthCS ay tumatanggap lamang ng CSV file para sa pag-upload ng dataset. "
    "Ito ay isang sadyang desisyon sa disenyo, hindi isang teknikal na limitasyon. "
    "Ang mga synthesis models — CTGAN at Gaussian Copula — ay parehong nangangailangan ng flat, tabular data: "
    "isang simpleng rows-at-columns na istruktura. "
    "Ang mga format tulad ng JSON o SQL ay maaaring hierarchical o multi-table, "
    "na mangangailangan ng karagdagang preprocessing at schema-flattening bago magsimula ang synthesis. "
    "Ang CSV ay din ang universal na export format sa mga spreadsheets (Excel, Google Sheets), "
    "databases (MySQL Workbench, pgAdmin), at data science tools (pandas, R). "
    "Ibig sabihin, ang mga user ay madaling makakapag-convert mula sa halos anumang source sa iisang hakbang. "
    "Inuna ng desisyon ang wastong pagtatayo ng synthesis pipeline kaysa sa pagsuporta sa maraming import format."
)

add_heading(doc, "12.2 Export Formats (5 Formats Supported)", 2)
section(doc,
    "EN: While import is restricted to CSV, SynthCS supports five export formats for the generated synthetic dataset. "
    "This asymmetry exists because export is simple serialization — the data is already a clean DataFrame, "
    "and it is just written out in a different format. Import is the harder direction because it involves "
    "trusting an unknown user file and normalizing it into a usable table.",
    "TL: Habang ang import ay limitado sa CSV, sinusuportahan ng SynthCS ang limang export format para sa "
    "nagawang synthetic dataset. Ang asymmetry na ito ay umiiral dahil ang export ay simpleng serialization — "
    "ang data ay isang malinis na DataFrame na, at isinusulat lamang ito sa ibang format. "
    "Ang import ang mas mahirap na direksyon dahil kailangan nitong pagkatiwalaan ang isang hindi kilalang file ng user "
    "at i-normalize ito sa isang magagamit na table."
)

export_formats = [
    ("CSV",   "Comma-separated values — universal tabular format",
               "Comma-separated values — universal na tabular format"),
    ("JSON",  "Array of objects — suitable for REST APIs and JavaScript apps",
               "Array ng mga objects — angkop para sa REST APIs at JavaScript apps"),
    ("JSONL", "One JSON object per line — optimized for streaming and LLM training data",
               "Isang JSON object bawat linya — optimized para sa streaming at LLM training data"),
    ("SQL",   "INSERT statements — ready to load directly into a relational database",
               "Mga INSERT statement — handa nang i-load direkta sa isang relational database"),
    ("Excel", ".xlsx workbook — for non-technical users and business stakeholders",
               ".xlsx workbook — para sa mga hindi teknikal na user at business stakeholders"),
]
for fmt, en_desc, tl_desc in export_formats:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"• {fmt}: ")
    r.bold = True; r.font.color.rgb = RGBColor(0x6d, 0x28, 0xd9)
    p.add_run(f"{en_desc}\n    TL: {tl_desc}")
    p.paragraph_format.space_after = Pt(4)

add_heading(doc, "12.3 Panel Talking Point", 2)
section(doc,
    "EN: \"We support CSV import because the synthesis models require flat, tabular data. "
    "Formats like JSON or SQL can be hierarchical or multi-table, which would need additional preprocessing "
    "before synthesis. CSV is also the universal export format for spreadsheets and databases, "
    "so users can easily convert from any source. We prioritized getting the synthesis pipeline right "
    "over supporting multiple import formats. "
    "Adding JSON or Excel import is straightforward — pandas already reads both natively — "
    "it is a planned enhancement, not a technical limitation.\"",
    "TL: \"Sinusuportahan namin ang CSV import dahil ang mga synthesis models ay nangangailangan ng flat, tabular data. "
    "Ang mga format tulad ng JSON o SQL ay maaaring hierarchical o multi-table, na mangangailangan ng karagdagang preprocessing "
    "bago ang synthesis. Ang CSV ay din ang universal na export format para sa mga spreadsheets at databases, "
    "kaya madaling makakapag-convert ang mga user mula sa anumang source. "
    "Inuna namin ang wastong pagtatayo ng synthesis pipeline kaysa sa pagsuporta sa maraming import format. "
    "Ang pagdaragdag ng JSON o Excel import ay straightforward — ang pandas ay nagbabasa ng pareho natively — "
    "ito ay isang planong enhancement, hindi isang teknikal na limitasyon.\""
)

# ── Save ──────────────────────────────────────────────────────────────────────
import os
out_path = os.path.join(os.path.dirname(__file__), "SynthCS_Backend_Reviewer.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
